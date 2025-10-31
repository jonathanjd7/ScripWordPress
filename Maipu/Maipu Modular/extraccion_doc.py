from docx import Document

def obtener_archivos_word_carpeta(carpeta):
    import os
    archivos = []
    for archivo in os.listdir(carpeta):
        if archivo.lower().endswith(".docx") and not archivo.startswith("~"):
            ruta_completa = os.path.join(carpeta, archivo)
            archivos.append(ruta_completa)
    return archivos

def extraer_titulo_doc(doc: Document):
    """Extrae el título del documento - VERSIÓN MEJORADA"""
    # Buscar con múltiples variaciones
    variaciones_titulo = [
        "titulo de la ficha",
        "título de la ficha",
        "titulo ficha",
        "título",
        "titulo:",
        "título:"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().lower()
        
        # Si encuentra alguna variación del marcador de título
        for variacion in variaciones_titulo:
            if variacion in texto:
                # Retornar el siguiente párrafo si existe
                if i + 1 < len(doc.paragraphs):
                    siguiente = doc.paragraphs[i + 1].text.strip()
                    if siguiente:
                        return siguiente
                # Si el título está en la misma línea después del marcador
                if ":" in para.text:
                    partes = para.text.split(":", 1)
                    if len(partes) > 1 and partes[1].strip():
                        return partes[1].strip()
    
    # Si no encontró con los marcadores, usar el primer párrafo no vacío
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto and len(texto) > 10:  # Al menos 10 caracteres
            return texto
    
    return ""

def frase_clave_doc(doc: Document):
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "frase clave objetivo" in texto.lower():
            buscar = True
    return ""

def titulo_seo_doc(doc: Document):
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "titulo seo" in texto.lower() or "ttulo seo" in texto.lower():
            buscar = True
    return ""

def meta_description_doc(doc: Document):
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "meta description" in texto.lower():
            buscar = True
    return ""

def leer_etiquetas_doc(doc: Document):
    buscar = False
    etiquetas = []
    palabras_clave = ["etiquetas", "etiqueta"]
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            etiquetas_divididas = texto.split(",")
            for etiq in etiquetas_divididas:
                etiq_limpia = etiq.strip()
                if etiq_limpia:
                    etiquetas.append(etiq_limpia)
            break
        if any(palabra in texto.lower() for palabra in palabras_clave):
            buscar = True
    return etiquetas

def leer_categorias_doc(doc: Document):
    """Extrae las categorías del documento y las divide por comas - VERSIÓN MEJORADA"""
    buscar = False
    categorias = []
    palabras_clave = ["categorias", "categoras", "categoria", "categori", "categoría"]
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            # Si no es otro encabezado de sección, procesar
            if not any(palabra in texto.lower() for palabra in ["etiqueta", "titulo de la ficha", "fin"]):
                # Dividir por comas si hay múltiples categorías
                if "," in texto:
                    cats_divididas = texto.split(",")
                    return [cat.strip() for cat in cats_divididas if cat.strip()]
                else:
                    # Si no hay comas, retornar como lista con un solo elemento
                    return [texto.strip()]
        if "categorías" in texto.lower() or "categorias" in texto.lower():
            buscar = True
    return []

def extraer_descripcion_con_formato_doc(doc: Document):
    """Extrae la descripción preservando el formato incluyendo hipervínculos - VERSIÓN COMPLETA"""
    descripcion_html = []
    capturar = False
    lista_abierta = False
    tipo_lista_actual = None  # 'ul' o 'ol'
    
    # Palabras clave que indican el fin de la descripción
    palabras_fin = [
        "frase clave objetivo",
        "título seo",
        "titulo seo",
        "meta description",
        "categorías",
        "categorias",
        "etiquetas",
        "fin descripción",
        "fin descripcion"
    ]
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        # Buscar inicio de descripción con múltiples variaciones
        if any(palabra in texto_lower for palabra in ["descripción", "descripcion", "descripción:", "descripcion:"]):
            capturar = True
            # Si la descripción está en la misma línea después de los dos puntos
            if ":" in texto and len(texto.split(":", 1)[1].strip()) > 0:
                texto_desc = texto.split(":", 1)[1].strip()
                if texto_desc:
                    descripcion_html.append(f"<p>{texto_desc}</p>")
            continue
        
        # Buscar fin de descripción con cualquiera de las palabras clave
        if capturar:
            debe_parar = any(palabra in texto_lower for palabra in palabras_fin)
            
            if debe_parar:
                # Cerrar lista si está abierta
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                break
            
        if capturar and texto:
            # Verificar el estilo del párrafo
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            
            # Detectar listas por estilo
            is_list_bullet = style_name in ["list paragraph", "viñeta", "list", "list bullet"]
            is_list_number = style_name in ["list number", "list 2", "list 3", "numbered list", "lista numerada"]
            
            # Detectar listas por numeración en XML (CLAVE PARA TU DOCUMENTO)
            tiene_numeracion = False
            try:
                if para._element.pPr is not None:
                    numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numPr is not None:
                        tiene_numeracion = True
                        # Verificar si es numerada (por defecto True si tiene numeración)
                        is_list_number = True
            except:
                pass
            
            is_list = is_list_bullet or is_list_number or tiene_numeracion
            
            # Detectar encabezados
            is_heading1 = style_name in ["heading 1", "título 1", "heading1"]
            is_heading2 = style_name in ["heading 2", "título 2", "heading2"]
            is_heading3 = style_name in ["heading 3", "título 3", "heading3"]
            
            # Los hipervínculos se procesarán solo como texto (sin <a href>)
            # Esto es más simple y evita problemas de extracción
            hyperlinks = {}
            
            # Procesar encabezados
            if is_heading1:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h2>{texto}</h2>")  # H2 para no competir con título principal
            elif is_heading2:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h3>{texto}</h3>")
            elif is_heading3:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h4>{texto}</h4>")
            elif is_list:
                # Determinar el tipo de lista (numerada u viñetas)
                tipo_lista_nueva = 'ol' if is_list_number else 'ul'
                
                # Si cambia el tipo de lista, cerrar la anterior y abrir nueva
                if lista_abierta and tipo_lista_actual != tipo_lista_nueva:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                
                # Abrir lista si no está abierta
                if not lista_abierta:
                    descripcion_html.append(f"<{tipo_lista_nueva}>")
                    lista_abierta = True
                    tipo_lista_actual = tipo_lista_nueva
                
                # Procesar contenido del elemento de lista con hipervínculos
                item_html = procesar_runs_con_formato_doc(para.runs, hyperlinks)
                
                if not item_html.strip():
                    item_html = texto
                
                descripcion_html.append(f"<li>{item_html}</li>")
            else:
                # Si era una lista y ahora no lo es, cerrar la lista
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                
                # Procesar párrafo normal con hipervínculos
                paragraph_html = procesar_runs_con_formato_doc(para.runs, hyperlinks)
            
            # Si después de procesar no hay contenido, usar el texto plano
            if not paragraph_html.strip():
                paragraph_html = texto
            
            # Envolver en párrafo HTML
            if paragraph_html.strip():
                descripcion_html.append(f"<p>{paragraph_html}</p>")
    
    # Cerrar lista si quedó abierta al final
    if lista_abierta:
        descripcion_html.append(f"</{tipo_lista_actual}>")
    
    return "\n".join(descripcion_html) if descripcion_html else ""

def procesar_runs_con_formato_doc(runs, hyperlinks=None):
    """Procesa los runs de un párrafo aplicando todos los formatos (negritas, cursivas, subrayado)"""
    if hyperlinks is None:
        hyperlinks = {}
    
    html = ""
    
    # Si no hay runs, retornar vacío
    if not runs:
        return html
    
    for run in runs:
        # Obtener el texto del run (puede ser vacío, con espacios, etc.)
        texto_run = run.text
        
        # Solo saltar si es None o string vacío (sin contar espacios)
        if texto_run is None or texto_run == "":
            continue
        
        # Los hipervínculos se muestran solo como texto (sin <a href>)
        # Aplicar formatos de fuente según las propiedades del run
        texto_formateado = texto_run
        
        # IMPORTANTE: Aplicar formatos solo si hay texto visible (no solo espacios)
        # pero mantener los espacios en el resultado
        if run.bold:
            texto_formateado = f"<strong>{texto_formateado}</strong>"
        if run.italic:
            texto_formateado = f"<em>{texto_formateado}</em>"
        if run.underline:
            texto_formateado = f"<u>{texto_formateado}</u>"
        
        html += texto_formateado
    
    return html
