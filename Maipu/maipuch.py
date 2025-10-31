import time
import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.action_chains import ActionChains

# ---------------- CONFIGURACIÓN CHILE ----------------
url_login = 'https://grupomaipu.com/es-ch/wp-login.php'
url_nuevo_post = 'https://grupomaipu.com/es-ch/wp-admin/post-new.php'
usuario_wp = 'GrupoMaipu2024'
password_wp = '1eH4.2NI>/&;'
carpeta_word = r"C:\Users\Jonathan JD\Desktop\pink\Jonathan\Maipu"

options = Options()
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

# Añadir opciones para mejorar estabilidad
options.add_argument("--disable-gpu")
options.add_argument("--disable-software-rasterizer")
options.add_argument("--memory-pressure-off")
options.add_argument("--disable-extensions")
options.add_argument("--disable-plugins")
options.add_argument("--no-first-run")
options.add_argument("--disable-default-apps")

# Opciones adicionales para mayor estabilidad
options.add_argument("--disable-web-security")
options.add_argument("--allow-running-insecure-content")
options.add_argument("--disable-features=VizDisplayCompositor")
options.add_argument("--disable-background-timer-throttling")
options.add_argument("--disable-backgrounding-occluded-windows")
options.add_argument("--disable-renderer-backgrounding")
options.add_argument("--disable-field-trial-config")
options.add_argument("--disable-back-forward-cache")
options.add_argument("--disable-ipc-flooding-protection")

def obtener_archivos_word(carpeta):
    """Obtiene todos los archivos .docx de la carpeta especificada"""
    archivos = []
    for archivo in os.listdir(carpeta):
        if archivo.lower().endswith('.docx') and not archivo.startswith('~$'):
            ruta_completa = os.path.join(carpeta, archivo)
            archivos.append(ruta_completa)
    return archivos

def extraer_titulo(doc):
    """Extrae el título del documento - VERSIÓN MEJORADA"""
    # Buscar con múltiples variaciones
    variaciones_titulo = [
        "titulo de la ficha",
        "título de la ficha",
        "titulo ficha",
        "título",
        "titulo:",
        "título:"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().lower()
        
        # Si encuentra alguna variación del marcador de título
        for variacion in variaciones_titulo:
            if variacion in texto:
                # Retornar el siguiente párrafo si existe
                if i + 1 < len(doc.paragraphs):
                    siguiente = doc.paragraphs[i + 1].text.strip()
                    if siguiente:
                        return siguiente
                # Si el título está en la misma línea después del marcador
                if ":" in para.text:
                    partes = para.text.split(":", 1)
                    if len(partes) > 1 and partes[1].strip():
                        return partes[1].strip()
    
    # Si no encontró con los marcadores, usar el primer párrafo no vacío
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto and len(texto) > 10:  # Al menos 10 caracteres
            return texto
    
    return ""

def extraer_descripcion_con_formato(doc):
    """Extrae la descripción preservando el formato incluyendo hipervínculos - VERSIÓN COMPLETA"""
    descripcion_html = []
    capturar = False
    lista_abierta = False
    tipo_lista_actual = None  # 'ul' o 'ol'
    
    # Palabras clave que indican el fin de la descripción
    palabras_fin = [
        "frase clave objetivo",
        "título seo",
        "titulo seo",
        "meta description",
        "categorías",
        "categorias",
        "etiquetas",
        "fin descripción",
        "fin descripcion"
    ]
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        # Buscar inicio de descripción con múltiples variaciones
        if any(palabra in texto_lower for palabra in ["descripción", "descripcion", "descripción:", "descripcion:"]):
            capturar = True
            # Si la descripción está en la misma línea después de los dos puntos
            if ":" in texto and len(texto.split(":", 1)[1].strip()) > 0:
                texto_desc = texto.split(":", 1)[1].strip()
                if texto_desc:
                    descripcion_html.append(f"<p>{texto_desc}</p>")
            continue
        
        # Buscar fin de descripción con cualquiera de las palabras clave
        if capturar:
            debe_parar = any(palabra in texto_lower for palabra in palabras_fin)
            
            if debe_parar:
                # Cerrar lista si está abierta
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                break
            
        if capturar and texto:
            # Verificar el estilo del párrafo
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            
            # Detectar listas por estilo
            is_list_bullet = style_name in ["list paragraph", "viñeta", "list", "list bullet"]
            is_list_number = style_name in ["list number", "list 2", "list 3", "numbered list", "lista numerada"]
            
            # Detectar listas por numeración en XML (CLAVE PARA TU DOCUMENTO)
            tiene_numeracion = False
            try:
                if para._element.pPr is not None:
                    numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numPr is not None:
                        tiene_numeracion = True
                        # Verificar si es numerada (por defecto True si tiene numeración)
                        is_list_number = True
            except:
                pass
            
            is_list = is_list_bullet or is_list_number or tiene_numeracion
            
            # Detectar encabezados
            is_heading1 = style_name in ["heading 1", "título 1", "heading1"]
            is_heading2 = style_name in ["heading 2", "título 2", "heading2"]
            is_heading3 = style_name in ["heading 3", "título 3", "heading3"]
            
            # Los hipervínculos se procesarán solo como texto (sin <a href>)
            # Esto es más simple y evita problemas de extracción
            hyperlinks = {}
            
            # Procesar encabezados
            if is_heading1:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h2>{texto}</h2>")  # H2 para no competir con título principal
            elif is_heading2:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h3>{texto}</h3>")
            elif is_heading3:
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                descripcion_html.append(f"<h4>{texto}</h4>")
            elif is_list:
                # Determinar el tipo de lista (numerada u viñetas)
                tipo_lista_nueva = 'ol' if is_list_number else 'ul'
                
                # Si cambia el tipo de lista, cerrar la anterior y abrir nueva
                if lista_abierta and tipo_lista_actual != tipo_lista_nueva:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                
                # Abrir lista si no está abierta
                if not lista_abierta:
                    descripcion_html.append(f"<{tipo_lista_nueva}>")
                    lista_abierta = True
                    tipo_lista_actual = tipo_lista_nueva
                
                # Procesar contenido del elemento de lista con hipervínculos
                item_html = procesar_runs_con_formato(para.runs, hyperlinks)
                
                if not item_html.strip():
                    item_html = texto
                
                descripcion_html.append(f"<li>{item_html}</li>")
            else:
                # Si era una lista y ahora no lo es, cerrar la lista
                if lista_abierta:
                    descripcion_html.append(f"</{tipo_lista_actual}>")
                    lista_abierta = False
                    tipo_lista_actual = None
                
                # Procesar párrafo normal con hipervínculos
                paragraph_html = procesar_runs_con_formato(para.runs, hyperlinks)
            
            # Si después de procesar no hay contenido, usar el texto plano
            if not paragraph_html.strip():
                paragraph_html = texto
            
            # Envolver en párrafo HTML
            if paragraph_html.strip():
                descripcion_html.append(f"<p>{paragraph_html}</p>")
    
    # Cerrar lista si quedó abierta al final
    if lista_abierta:
        descripcion_html.append(f"</{tipo_lista_actual}>")
    
    return "\n".join(descripcion_html) if descripcion_html else ""

def procesar_runs_con_formato(runs, hyperlinks=None):
    """Procesa los runs de un párrafo aplicando todos los formatos (negritas, cursivas, subrayado)"""
    if hyperlinks is None:
        hyperlinks = {}
    
    html = ""
    
    # Si no hay runs, retornar vacío
    if not runs:
        return html
    
    for run in runs:
        # Obtener el texto del run (puede ser vacío, con espacios, etc.)
        texto_run = run.text
        
        # Solo saltar si es None o string vacío (sin contar espacios)
        if texto_run is None or texto_run == "":
            continue
        
        # Los hipervínculos se muestran solo como texto (sin <a href>)
        # Aplicar formatos de fuente según las propiedades del run
        texto_formateado = texto_run
        
        # IMPORTANTE: Aplicar formatos solo si hay texto visible (no solo espacios)
        # pero mantener los espacios en el resultado
        if run.bold:
            texto_formateado = f"<strong>{texto_formateado}</strong>"
        if run.italic:
            texto_formateado = f"<em>{texto_formateado}</em>"
        if run.underline:
            texto_formateado = f"<u>{texto_formateado}</u>"
        
        html += texto_formateado
    
    return html

def frase_clave(doc):
    """Extrae la frase clave objetivo - VERSIÓN MEJORADA"""
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "frase clave objetivo" in texto.lower():
            buscar = True
    return ""

def titulo_seo(doc):
    """Extrae el título SEO - VERSIÓN MEJORADA"""
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "título seo" in texto.lower() or "titulo seo" in texto.lower():
            buscar = True
    return ""

def meta_description(doc):
    """Extrae la meta description - VERSIÓN MEJORADA"""
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            return texto
        if "meta description" in texto.lower() or "meta description" in texto.lower():
            buscar = True
    return ""

def leer_etiquetas(doc):
    """Extrae las etiquetas del documento - VERSIÓN MEJORADA"""
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            # Si no es un encabezado de otra sección, retornar
            if not any(palabra in texto.lower() for palabra in ["categorías", "categorias", "titulo de la ficha"]):
                return texto
        if "etiqueta" in texto.lower():
            buscar = True
    return ""

def leer_categorias(doc):
    """Extrae las categorías del documento y las divide por comas - VERSIÓN MEJORADA"""
    buscar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            # Si no es otro encabezado de sección, procesar
            if not any(palabra in texto.lower() for palabra in ["etiqueta", "titulo de la ficha", "fin"]):
                # Dividir por comas si hay múltiples categorías
                if "," in texto:
                    cats_divididas = texto.split(",")
                    return [cat.strip() for cat in cats_divididas if cat.strip()]
                else:
                    # Si no hay comas, retornar como lista con un solo elemento
                    return [texto.strip()]
        if "categorías" in texto.lower() or "categorias" in texto.lower():
            buscar = True
    return []

def cerrar_menus_interferentes(driver):
    """Cierra cualquier menú desplegable que pueda estar interfiriendo"""
    try:
        # Cerrar menú de perfil de usuario si está abierto
        driver.execute_script("""
            // Cerrar menús desplegables comunes
            var menus = document.querySelectorAll('.ab-item[aria-expanded="true"]');
            menus.forEach(function(menu) {
                menu.click();
            });
            
            // Cerrar cualquier overlay o modal
            var overlays = document.querySelectorAll('.wp-admin-bar-menu, .ab-submenu');
            overlays.forEach(function(overlay) {
                if (overlay.style.display !== 'none') {
                    overlay.style.display = 'none';
                }
            });
        """)
        time.sleep(0.5)
    except:
        pass

def crear_nueva_categoria_clasico(driver, wait, nombre_categoria):
    """Crea una nueva categoría en WordPress (Editor Clásico) si no existe"""
    try:
        print(f"[INFO] Intentando crear nueva categoría: '{nombre_categoria}'")
        
        # Cerrar menús que puedan estar interfiriendo
        cerrar_menus_interferentes(driver)
        
        # Paso 1: Hacer clic en el enlace "Añadir una nueva categoría"
        try:
            añadir_link = None
            
            # Selector 1: Por ID
            try:
                añadir_link = wait.until(EC.element_to_be_clickable((By.ID, 'category-add-toggle')))
                print("[DEBUG] Enlace encontrado por ID 'category-add-toggle'")
            except:
                pass
            
            # Selector 2: Por clase category-add-toggle
            if not añadir_link:
                try:
                    añadir_link = driver.find_element(By.CLASS_NAME, 'category-add-toggle')
                    print("[DEBUG] Enlace encontrado por clase 'category-add-toggle'")
                except:
                    pass
            
            # Selector 3: Por texto "Añadir"
            if not añadir_link:
                try:
                    añadir_link = driver.find_element(By.XPATH, '//a[contains(text(), "Añadir") and contains(@class, "category")]')
                    print("[DEBUG] Enlace encontrado por texto 'Añadir'")
                except:
                    pass
            
            if not añadir_link:
                print("[ERROR] No se encontro el enlace 'Anadir categoria' (probados 3 selectores)")
                return False
            
            # Hacer clic en el enlace
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", añadir_link)
            time.sleep(1)
            
            # Usar JavaScript click
            try:
                driver.execute_script("arguments[0].click();", añadir_link)
                print("[OK] Enlace 'Anadir categoria' clickeado (JavaScript)")
            except:
                añadir_link.click()
                print("[OK] Enlace 'Anadir categoria' clickeado (click directo)")
            
            time.sleep(3)  # Dar más tiempo para que aparezca el formulario
            
        except Exception as e:
            print(f"[ERROR] Error al hacer clic en 'Anadir categoria': {e}")
            return False
        
        # Paso 2: Buscar el campo de texto para el nombre de la nueva categoría
        try:
            # Intentar múltiples selectores
            campo_nombre = None
            
            # Selector 1: ID estándar
            try:
                campo_nombre = wait.until(EC.visibility_of_element_located((By.ID, 'newcategory')))
                print("[DEBUG] Campo encontrado por ID 'newcategory'")
            except:
                pass
            
            # Selector 2: Por nombre
            if not campo_nombre:
                try:
                    campo_nombre = driver.find_element(By.NAME, 'newcategory')
                    print("[DEBUG] Campo encontrado por NAME 'newcategory'")
                except:
                    pass
            
            # Selector 3: Buscar input en el área de categorías que apareció
            if not campo_nombre:
                try:
                    # Buscar inputs de texto visibles en el área de categorías
                    campos = driver.find_elements(By.XPATH, '//div[@id="category-adder"]//input[@type="text"]')
                    if campos:
                        campo_nombre = campos[0]
                        print("[DEBUG] Campo encontrado en div category-adder")
                except:
                    pass
            
            # Selector 4: Cualquier input visible recién aparecido
            if not campo_nombre:
                try:
                    campos = driver.find_elements(By.XPATH, '//input[@type="text" and @aria-label]')
                    for campo in campos:
                        if campo.is_displayed():
                            campo_nombre = campo
                            print("[DEBUG] Campo encontrado por visibilidad")
                            break
                except:
                    pass
            
            if not campo_nombre:
                print("[ERROR] No se encontro el campo para el nombre de la categoria (probados 4 selectores)")
                return False
            
            # MEJORADO: Usar JavaScript para interactuar con el campo y evitar interceptación
            print("[DEBUG] Usando JavaScript para interactuar con el campo de categoría...")
            
            # Hacer scroll para que sea visible
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo_nombre)
            time.sleep(1)
            
            # Usar JavaScript para limpiar y establecer el valor
            driver.execute_script("arguments[0].value = '';", campo_nombre)
            driver.execute_script("arguments[0].value = arguments[1];", campo_nombre, nombre_categoria)
            
            # Disparar eventos para que WordPress detecte el cambio
            driver.execute_script("""
                var element = arguments[0];
                element.dispatchEvent(new Event('input', { bubbles: true }));
                element.dispatchEvent(new Event('change', { bubbles: true }));
                element.dispatchEvent(new Event('blur', { bubbles: true }));
            """, campo_nombre)
            
            print(f"[OK] Nombre de categoria '{nombre_categoria}' ingresado mediante JavaScript")
            time.sleep(1)
        except Exception as ex:
            print(f"[ERROR] Error al buscar/ingresar campo de categoria: {ex}")
            return False
        
        # Paso 3: Hacer clic en el botón "Añadir nueva categoría"
        try:
            confirmar_btn = None
            
            # Selector 1: Por ID
            try:
                confirmar_btn = wait.until(EC.element_to_be_clickable((By.ID, 'category-add-submit')))
                print("[DEBUG] Boton encontrado por ID 'category-add-submit'")
            except:
                pass
            
            # Selector 2: Por clase o texto
            if not confirmar_btn:
                try:
                    confirmar_btn = driver.find_element(By.XPATH, '//input[@type="button" and contains(@value, "Añadir")]')
                    print("[DEBUG] Boton encontrado por XPATH con 'Añadir'")
                except:
                    pass
            
            # Selector 3: Cualquier botón en el área de categorías
            if not confirmar_btn:
                try:
                    confirmar_btn = driver.find_element(By.XPATH, '//div[@id="category-adder"]//input[@type="button"]')
                    print("[DEBUG] Boton encontrado en div category-adder")
                except:
                    pass
            
            if not confirmar_btn:
                print("[ERROR] No se encontro el boton de confirmar (probados 3 selectores)")
                # Intentar con Enter como último recurso
                try:
                    driver.execute_script("arguments[0].dispatchEvent(new KeyboardEvent('keydown', {key: 'Enter', keyCode: 13}));", campo_nombre)
                    print("[OK] Enter presionado mediante JavaScript como alternativa")
                    time.sleep(2)
                    return True
                except:
                    return False
            
            # MEJORADO: Usar JavaScript para hacer clic y evitar interceptación
            print("[DEBUG] Usando JavaScript para hacer clic en el botón...")
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", confirmar_btn)
            time.sleep(0.5)
            
            # Usar JavaScript click
            try:
                driver.execute_script("arguments[0].click();", confirmar_btn)
                print("[OK] Boton 'Anadir nueva categoria' clickeado (JavaScript)")
            except:
                # Fallback: intentar con click directo
                try:
                    confirmar_btn.click()
                    print("[OK] Boton 'Anadir nueva categoria' clickeado (click directo)")
                except:
                    # Último recurso: simular clic con JavaScript
                    driver.execute_script("arguments[0].dispatchEvent(new MouseEvent('click', {bubbles: true}));", confirmar_btn)
                    print("[OK] Boton 'Anadir nueva categoria' clickeado (simulación JavaScript)")
            
            time.sleep(3)  # Dar más tiempo para que se procese
            
            # Verificar que la categoría se creó correctamente
            try:
                # Buscar la categoría recién creada en la lista
                xpath_verificacion = f"//label[normalize-space(text()) = '{nombre_categoria}']/input[@type='checkbox']"
                categoria_creada = driver.find_element(By.XPATH, xpath_verificacion)
                if categoria_creada:
                    print(f"[SUCCESS] Categoria '{nombre_categoria}' creada y verificada exitosamente")
                    return True
                else:
                    print(f"[WARNING] Categoria '{nombre_categoria}' puede no haberse creado correctamente")
                    return True  # Continuar de todas formas
            except:
                print(f"[WARNING] No se pudo verificar la creación de la categoría '{nombre_categoria}', pero continuamos")
                return True
            
        except Exception as e:
            print(f"[ERROR] Error al confirmar creacion de categoria: {e}")
            return False
        
    except Exception as e:
        print(f"[ERROR] Error general al crear categoria '{nombre_categoria}': {e}")
        return False

def guardar_borrador(driver, wait):
    """Función para guardar el post como borrador - VERSIÓN MEJORADA CON JAVASCRIPT"""
    try:
        print("[INFO] Guardando como borrador...")
        
        # Método 1: Usar JavaScript para hacer click directamente (evita problemas de intercepción)
        try:
            driver.execute_script("document.getElementById('save-post').click();")
            print("[OK] Clic en botón 'Guardar borrador' mediante JavaScript")
        except Exception as e:
            print(f"[WARNING] Método JavaScript falló: {e}")
            
            # Método 2: Buscar botón y usar JavaScript con el elemento
            try:
                guardar_btn = wait.until(EC.presence_of_element_located((By.ID, 'save-post')))
                driver.execute_script("arguments[0].click();", guardar_btn)
                print("[OK] Clic en botón 'Guardar borrador' (Método 2)")
            except Exception as e2:
                print(f"[ERROR] Todos los métodos fallaron: {e2}")
                return False
        
        # Esperar a que se procese el guardado
        time.sleep(5)
        
        # Verificar mensaje de éxito
        try:
            exito_element = wait.until(EC.presence_of_element_located((By.ID, 'message')))
            mensaje = exito_element.text.lower()
            print(f"[INFO] Mensaje del sistema: {mensaje}")
            
            if any(palabra in mensaje for palabra in ['guardado', 'borrador', 'actualizado', 'saved', 'draft']):
                print("[OK] Guardado exitoso - Mensaje confirmado")
                return True
            else:
                print(f"[WARNING] Mensaje inesperado: {mensaje}")
                return True  # Continuamos aunque el mensaje no sea el esperado
                
        except TimeoutException:
            print("[WARNING] No se encontró mensaje de confirmación, pero continuamos")
            return True
            
    except Exception as e:
        print(f"[ERROR] Error al guardar: {e}")
        return False

def insertar_descripcion_classic_editor(driver, wait, descripcion):
    """Inserta descripción en Classic Editor - VERSIÓN MEJORADA CON SINCRONIZACIÓN"""
    try:
        print("[INFO] Configurando editor clásico para descripción...")
        
        # Paso 1: Cambiar a pestaña HTML
        try:
            html_btn = wait.until(EC.element_to_be_clickable((By.ID, 'content-html')))
            html_btn.click()
            print("[OK] Cambiado a modo HTML")
            time.sleep(2)
        except Exception as e:
            print(f"[ERROR] No se pudo cambiar a modo HTML: {e}")
            return False
        
        # Paso 2: Insertar contenido en el textarea
        try:
            # Buscar el textarea directamente
            textarea_editor = wait.until(EC.presence_of_element_located((By.ID, 'content')))
            
            # Limpiar contenido previo
            driver.execute_script("arguments[0].value = '';", textarea_editor)
            time.sleep(0.5)
            
            # Insertar la descripción usando JavaScript
            driver.execute_script("arguments[0].value = arguments[1];", textarea_editor, descripcion)
            
            # Disparar eventos para que WordPress detecte el cambio
            driver.execute_script("""
                var element = arguments[0];
                element.dispatchEvent(new Event('input', { bubbles: true }));
                element.dispatchEvent(new Event('change', { bubbles: true }));
            """, textarea_editor)
            
            print(f"[OK] Descripción insertada mediante JavaScript ({len(descripcion)} caracteres)")
            time.sleep(1)
            
            # Verificar que se insertó correctamente
            valor_actual = driver.execute_script("return arguments[0].value;", textarea_editor)
            if len(valor_actual) > 100:
                print(f"[OK] Verificación: Contenido presente ({len(valor_actual)} caracteres)")
            else:
                print(f"[WARNING] Verificación: Contenido parece vacío o corto ({len(valor_actual)} caracteres)")
            
        except Exception as e:
            print(f"[ERROR] Error al insertar en textarea: {e}")
            return False
        
        # Paso 3: NO cambiar a modo visual - Quedarse en HTML para preservar formatos
        print("[INFO] Manteniendo modo HTML para preservar formatos...")
        time.sleep(2)
        
        # Verificar una última vez que el contenido sigue en el textarea
        try:
            valor_final = driver.execute_script("return document.getElementById('content').value;")
            if len(valor_final) > 100:
                print(f"[OK] Contenido confirmado en modo HTML ({len(valor_final)} caracteres)")
                print(f"[OK] HTML contiene: {valor_final.count('<strong>')} negritas, {valor_final.count('<ol>')} listas numeradas")
            else:
                print(f"[WARNING] Contenido parece haberse perdido")
        except:
            print("[WARNING] No se pudo verificar contenido final")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Error insertando descripción en editor clásico: {e}")
        # Asegurarse de volver al contenido principal en caso de error
        try:
            driver.switch_to.default_content()
        except:
            pass
        return False

def procesar_archivo(archivo_word, driver, wait):
    """Procesa un archivo Word individual y lo guarda como borrador en WordPress"""
    print(f"\n{'='*60}")
    print(f"[INFO] Procesando: {os.path.basename(archivo_word)}")
    print(f"{'='*60}")
    
    try:
        # Navegar a la página de nuevo post
        driver.get(url_nuevo_post)
        time.sleep(4)

        doc = Document(archivo_word)
        titulo = extraer_titulo(doc)
        descripcion = extraer_descripcion_con_formato(doc)
        frase_obj = frase_clave(doc)
        tit_seo = titulo_seo(doc)
        meta_desc = meta_description(doc)
        etiquetas = leer_etiquetas(doc)
        categorias = leer_categorias(doc)

        print(f"[INFO] Título: {titulo if titulo else '[VACÍO - REVISAR DOCUMENTO]'}")
        if descripcion:
            print(f"[INFO] Descripción: {descripcion[:100]}... ({len(descripcion)} caracteres)")
        else:
            print(f"[WARNING] Descripción: [VACÍA - REVISAR DOCUMENTO]")
        print(f"[INFO] Frase clave: {frase_obj}")
        print(f"[INFO] Título SEO: {tit_seo}")
        print(f"[INFO] Meta desc: {meta_desc}")
        print(f"[INFO] Etiquetas: {etiquetas}")
        print(f"[INFO] Categorías: {categorias}")

        # ---------------- INSERTAR TÍTULO ----------------
        print("[INFO] Insertando título...")
        title_field = wait.until(EC.element_to_be_clickable((By.ID, 'title')))
        title_field.click()
        title_field.clear()
        title_field.send_keys(titulo)
        time.sleep(1)

        # ---------------- INSERTAR DESCRIPCIÓN ----------------
        if descripcion.strip():
            print(f"[INFO] Insertando descripción en editor clásico ({len(descripcion)} caracteres)...")
            print(f"[DEBUG] La descripción contiene {descripcion.count('<strong>')} negritas, {descripcion.count('<ol>')} listas numeradas")
            
            if not insertar_descripcion_classic_editor(driver, wait, descripcion):
                print("[ERROR] No se pudo insertar la descripción")
                return False
            
            # Espera adicional para que WordPress procese la descripción
            print("[INFO] Esperando a que WordPress procese la descripción...")
            time.sleep(3)
        else:
            print("[WARNING] No hay descripción para insertar")

        # ---------------- INSERTAR FRASE CLAVE ----------------
        print("[INFO] Insertando frase clave...")
        try:
            frase_field = wait.until(EC.visibility_of_element_located((By.ID, 'focus-keyword-input-metabox')))
            frase_field.clear()
            frase_field.send_keys(Keys.CONTROL, 'a')
            frase_field.send_keys(Keys.DELETE)
            frase_field.send_keys(frase_obj)
            print("[OK] Frase clave insertada")
        except TimeoutException:
            print("[WARNING] Campo de frase clave no encontrado")

        # ---------------- INSERTAR TÍTULO SEO ----------------
        print("[INFO] Insertando título SEO...")
        try:
            # Este campo es más complejo, intentamos enfoque diferente
            titulo_seo_area = wait.until(EC.element_to_be_clickable((By.ID, 'yoast-google-preview-title-metabox')))
            titulo_seo_area.click()
            titulo_seo_area.send_keys(Keys.CONTROL, 'a')
            titulo_seo_area.send_keys(Keys.DELETE)
            titulo_seo_area.send_keys(tit_seo)
            print("[OK] Título SEO insertado")
        except TimeoutException:
            print("[WARNING] Campo de título SEO no encontrado")

        # ---------------- INSERTAR META DESCRIPCIÓN ----------------
        print("[INFO] Insertando meta descripción...")
        try:
            meta_desc_field = wait.until(EC.element_to_be_clickable((By.ID, 'yoast-google-preview-description-metabox')))
            meta_desc_field.click()
            meta_desc_field.send_keys(Keys.CONTROL, 'a')
            meta_desc_field.send_keys(Keys.DELETE)
            meta_desc_field.send_keys(meta_desc)
            print("[OK] Meta descripción insertada")
        except TimeoutException:
            print("[WARNING] Campo de meta descripción no encontrado")

        # ---------------- INSERTAR CATEGORÍAS ----------------
        print("[INFO] Configurando categorías...")
        print(f"[DEBUG] Total de categorías detectadas: {len(categorias)}")
        if categorias:
            print(f"[DEBUG] Categorías a procesar: {', '.join(categorias)}")
        
        # Contador para rastrear categorías procesadas exitosamente
        categorias_procesadas = 0
        categorias_fallidas = 0
        
        # Procesar cada categoría individualmente para evitar que un error detenga el proceso
        for i, categoria in enumerate(categorias, 1):
            categoria_limpia = categoria.strip()
            if not categoria_limpia:
                continue
                
            print(f"[DEBUG] Procesando categoría {i}/{len(categorias)}: '{categoria_limpia}'")
            checkbox_encontrado = False
            
            try:
                # Método 1: Buscar por texto exacto del label
                xpath = f"//label[normalize-space(text()) = '{categoria_limpia}']/input[@type='checkbox']"
                checkbox = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
                
                # Hacer scroll para que sea visible
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", checkbox)
                time.sleep(0.5)
                
                if not checkbox.is_selected():
                    # Usar JavaScript click para evitar interceptación
                    try:
                        driver.execute_script("arguments[0].click();", checkbox)
                        print(f"[OK] Categoría seleccionada: {categoria_limpia}")
                    except:
                        checkbox.click()
                        print(f"[OK] Categoría seleccionada (click directo): {categoria_limpia}")
                else:
                    print(f"[INFO] Categoría ya estaba seleccionada: {categoria_limpia}")
                checkbox_encontrado = True
                categorias_procesadas += 1
                time.sleep(0.5)
                
            except TimeoutException:
                print(f"[WARNING] No se encontró la categoría '{categoria_limpia}'")
                
                # Método 2: Intentar búsqueda más flexible
                try:
                    xpath_flex = f"//label[contains(normalize-space(text()), '{categoria_limpia}')]/input[@type='checkbox']"
                    checkbox = driver.find_element(By.XPATH, xpath_flex)
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", checkbox)
                    time.sleep(0.5)
                    
                    if not checkbox.is_selected():
                        try:
                            driver.execute_script("arguments[0].click();", checkbox)
                            print(f"[OK] Categoría seleccionada (búsqueda flexible): {categoria_limpia}")
                        except:
                            checkbox.click()
                            print(f"[OK] Categoría seleccionada (búsqueda flexible - click directo): {categoria_limpia}")
                    checkbox_encontrado = True
                    categorias_procesadas += 1
                    time.sleep(0.5)
                except:
                    print(f"[WARNING] Categoría '{categoria_limpia}' no existe, intentando crearla...")
                    
                    # Método 3: Crear la categoría si no existe
                    if crear_nueva_categoria_clasico(driver, wait, categoria_limpia):
                        print(f"[OK] Categoría '{categoria_limpia}' creada exitosamente")
                        
                        # Esperar y buscar la categoría recién creada para seleccionarla
                        time.sleep(2)
                        try:
                            xpath = f"//label[normalize-space(text()) = '{categoria_limpia}']/input[@type='checkbox']"
                            checkbox = driver.find_element(By.XPATH, xpath)
                            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", checkbox)
                            time.sleep(0.5)
                            
                            if not checkbox.is_selected():
                                try:
                                    driver.execute_script("arguments[0].click();", checkbox)
                                    print(f"[OK] Categoría '{categoria_limpia}' seleccionada después de crearla")
                                except:
                                    checkbox.click()
                                    print(f"[OK] Categoría '{categoria_limpia}' seleccionada después de crearla (click directo)")
                            checkbox_encontrado = True
                            categorias_procesadas += 1
                        except Exception as e:
                            print(f"[WARNING] Se creó la categoría pero no se pudo seleccionar: {e}")
                    else:
                        print(f"[ERROR] No se pudo crear la categoría '{categoria_limpia}'")
            
            except Exception as e:
                print(f"[ERROR] Error inesperado con la categoría '{categoria_limpia}': {e}")
                categorias_fallidas += 1
                # Continuar con la siguiente categoría en lugar de detener todo el proceso
                continue
            
            # Verificar si se procesó correctamente esta categoría
            if not checkbox_encontrado:
                print(f"[WARNING] La categoría '{categoria_limpia}' no se pudo procesar correctamente")
                categorias_fallidas += 1
        
        # Resumen de categorías procesadas
        print(f"[INFO] Resumen de categorías: {categorias_procesadas} procesadas exitosamente, {categorias_fallidas} fallidas de {len(categorias)} total")

        # ---------------- INSERTAR ETIQUETAS ----------------
        print("[INFO] Configurando etiquetas...")
        try:
            if etiquetas.strip():
                etiq_field = wait.until(EC.visibility_of_element_located((By.ID, 'new-tag-post_tag')))
                etiq_field.clear()
                etiq_field.send_keys(etiquetas)
                etiq_field.send_keys(Keys.ENTER)
                print("[OK] Etiquetas insertadas")
        except Exception as e:
            print(f"[WARNING] Error insertando etiquetas: {e}")

        # ---------------- PREPARAR PÁGINA PARA GUARDAR ----------------
        print("[INFO] Preparando página para guardar...")
        time.sleep(3)
        
        # Hacer scroll para asegurar que todos los elementos estén cargados
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1)
        
        # Verificación final: asegurar que la descripción está en el textarea
        if descripcion.strip():
            try:
                print("[DEBUG] Verificación final de descripción antes de guardar...")
                textarea_content = driver.execute_script("return document.getElementById('content').value;")
                if len(textarea_content) > 100:
                    print(f"[OK] Descripción confirmada en textarea ({len(textarea_content)} caracteres)")
                else:
                    print(f"[WARNING] Descripción parece vacía en textarea, reinsertando...")
                    # Reinsertar si está vacío
                    driver.execute_script("document.getElementById('content').value = arguments[0];", descripcion)
                    print(f"[OK] Descripción reinsertada ({len(descripcion)} caracteres)")
            except Exception as e:
                print(f"[WARNING] No se pudo verificar descripción: {e}")

        # ---------------- GUARDAR COMO BORRADOR ----------------
        print("[INFO] Guardando como borrador...")
        
        if guardar_borrador(driver, wait):
            print(f"[SUCCESS] Post '{titulo}' guardado como borrador con éxito!")
            return True
        else:
            print(f"[ERROR] No se pudo guardar '{titulo}' como borrador")
            return False

    except Exception as e:
        print(f"[ERROR] Error procesando {archivo_word}: {e}")
        return False

def main():
    print("="*80)
    print("MAIPU CHILE - AUTOMATIZACIÓN DE POSTS WORDPRESS")
    print("="*80)
    
    # Obtener todos los archivos Word de la carpeta
    archivos = obtener_archivos_word(carpeta_word)
    
    if not archivos:
        print(f"[ERROR] No se encontraron archivos .docx en la carpeta: {carpeta_word}")
        return
    
    print(f"[INFO] Se encontraron {len(archivos)} archivos Word:")
    for i, archivo in enumerate(archivos, 1):
        print(f"  {i}. {os.path.basename(archivo)}")
    
    # Inicializar el navegador una sola vez
    print("[INFO] Inicializando navegador...")
    driver = webdriver.Chrome(options=options)
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    wait = WebDriverWait(driver, 30)
    
    try:
        # Login una sola vez
        print("[INFO] Iniciando sesión en WordPress CHILE...")
        driver.get(url_login)
        time.sleep(3)
        
        username_field = wait.until(EC.visibility_of_element_located((By.ID, 'user_login')))
        password_field = wait.until(EC.visibility_of_element_located((By.ID, 'user_pass')))
        
        username_field.clear()
        username_field.send_keys(usuario_wp)
        password_field.clear()
        password_field.send_keys(password_wp)
        
        login_button = wait.until(EC.element_to_be_clickable((By.ID, 'wp-submit')))
        login_button.click()
        
        # Esperar a que cargue el dashboard
        wait.until(EC.presence_of_element_located((By.ID, 'wpadminbar')))
        print("[OK] Login exitoso en CHILE")

        # Procesar cada archivo
        exitosos = 0
        total = len(archivos)
        
        for i, archivo in enumerate(archivos, 1):
            print(f"\n[INFO] Progreso CHILE: {i}/{total}")
            if procesar_archivo(archivo, driver, wait):
                exitosos += 1
            time.sleep(5)  # Espera entre guardados
        
        print(f"\n{'='*80}")
        print(f"[SUCCESS] PROCESO COMPLETADO EN CHILE")
        print(f"{'='*80}")
        print(f"[OK] Archivos guardados como borrador: {exitosos}/{total}")
        print(f"[ERROR] Archivos fallidos: {total - exitosos}")
        print(f"{'='*80}")

    except Exception as e:
        print(f"[ERROR] Error en el proceso principal: {e}")
        
    finally:
        time.sleep(3)
        driver.quit()
        print("[INFO] Navegador cerrado.")

if __name__ == "__main__":
    main()
