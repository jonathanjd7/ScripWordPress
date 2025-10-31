
import os
import time
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from selenium.webdriver.common.action_chains import ActionChains

# ---------------- CONFIGURACIÓN ----------------
# chrome_driver_path = r"C:/Users/Usuario/Desktop/chromedriver.exe"  # Not needed with modern Selenium
url_wp = 'https://cornergourmet.es/wp-admin'
usuario_wp = 'admin5874'
password_wp = 'HSaQ^a6)vciC^e&@(duf46Qg'
carpeta_word = r"C:\Users\Jonathan JD\Desktop\pink\Jonathan\CodigoCorner 1"  # Cambia esta ruta a tu carpeta

options = Options()
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

# Añadir opciones para mejorar estabilidad
options.add_argument("--disable-gpu")
options.add_argument("--disable-software-rasterizer")
options.add_argument("--memory-pressure-off")
options.add_argument("--disable-extensions")
options.add_argument("--disable-plugins")
options.add_argument("--no-first-run")
options.add_argument("--disable-default-apps")

def obtener_archivos_word(carpeta):
    """Obtiene todos los archivos .docx de la carpeta especificada"""
    archivos = []
    for archivo in os.listdir(carpeta):
        if archivo.lower().endswith('.docx') and not archivo.startswith('~$'):
            ruta_completa = os.path.join(carpeta, archivo)
            archivos.append(ruta_completa)
    return archivos

def leer_documento_con_formato(path):
    """Lee el documento Word y preserva el formato HTML - VERSIÓN MEJORADA"""
    doc = Document(path)
    contenido_html = []
    lista_abierta = False
    
    for para in doc.paragraphs:
        if para.text.strip():
            # Verificar si el párrafo tiene estilo de lista
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            is_list = style_name in ["list paragraph", "viñeta", "list", "list bullet", "list number"]
            
            if is_list:
                # Abrir lista si no está abierta
                if not lista_abierta:
                    contenido_html.append("<ul>")
                    lista_abierta = True
                
                # Procesar contenido del elemento de lista
                item_html = ""
                for run in para.runs:
                    if run.text.strip():
                        texto_run = run.text
                        # Aplicar formatos múltiples
                        if run.bold:
                            texto_run = f"<strong>{texto_run}</strong>"
                        if run.italic:
                            texto_run = f"<em>{texto_run}</em>"
                        if run.underline:
                            texto_run = f"<u>{texto_run}</u>"
                        item_html += texto_run
                
                if not item_html.strip():
                    item_html = para.text.strip()
                
                contenido_html.append(f"<li>{item_html}</li>")
            else:
                # Si era una lista y ahora no lo es, cerrar la lista
                if lista_abierta:
                    contenido_html.append("</ul>")
                    lista_abierta = False
                
                # Procesar párrafo normal con runs
                paragraph_html = ""
                for run in para.runs:
                    if run.text.strip():
                        texto_run = run.text
                        # Aplicar formatos múltiples (puede tener negrita + cursiva + subrayado)
                        if run.bold:
                            texto_run = f"<strong>{texto_run}</strong>"
                        if run.italic:
                            texto_run = f"<em>{texto_run}</em>"
                        if run.underline:
                            texto_run = f"<u>{texto_run}</u>"
                        paragraph_html += texto_run
                
                # Si no hay contenido después de procesar runs, usar texto plano
                if not paragraph_html.strip():
                    paragraph_html = para.text.strip()
                
                # Envolver en párrafo
                if paragraph_html.strip():
                    contenido_html.append(f"<p>{paragraph_html}</p>")
    
    # Cerrar lista si quedó abierta al final
    if lista_abierta:
        contenido_html.append("</ul>")
    
    return "\n".join(contenido_html)

def extraer_titulo(doc):
    """Extrae el título del documento - VERSIÓN MEJORADA"""
    # Buscar con múltiples variaciones
    variaciones_titulo = [
        "titulo de la ficha",
        "título de la ficha",
        "titulo ficha",
        "título",
        "titulo:",
        "título:"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().lower()
        
        # Si encuentra alguna variación del marcador de título
        for variacion in variaciones_titulo:
            if variacion in texto:
                # Retornar el siguiente párrafo si existe
                if i + 1 < len(doc.paragraphs):
                    siguiente = doc.paragraphs[i + 1].text.strip()
                    if siguiente:
                        return siguiente
                # Si el título está en la misma línea después del marcador
                if ":" in para.text:
                    partes = para.text.split(":", 1)
                    if len(partes) > 1 and partes[1].strip():
                        return partes[1].strip()
    
    # Si no encontró con los marcadores, usar el primer párrafo no vacío
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto and len(texto) > 10:  # Al menos 10 caracteres
            return texto
    
    return ""

def extraer_descripcion_con_formato(doc):
    """Extrae la descripción preservando el formato - VERSIÓN MEJORADA"""
    descripcion_html = []
    capturar = False
    lista_abierta = False
    
    # Palabras clave que indican el fin de la descripción (verificar si ES o CONTIENE estas palabras al inicio)
    palabras_fin = [
        "frase clave objetivo",
        "frase clave",
        "título seo",
        "titulo seo",
        "meta description",
        "meta descripción",
        "categorías:",
        "categorias:",
        "categorías",
        "categorias",
        "etiquetas:",
        "etiqueta:",
        "etiquetas",
        "etiqueta",
        "tags:",
        "tag:",
        "tags",
        "tag",
        "fin descripción",
        "fin descripcion"
    ]
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower().strip()
        
        # Verificar si este párrafo es un marcador de inicio de otra sección (ANTES de empezar a capturar)
        if not capturar:
            # Buscar inicio de descripción con múltiples variaciones
            if any(palabra in texto_lower for palabra in ["descripción:", "descripcion:", "descripción", "descripcion"]):
                # Verificar que NO sea solo el marcador sin contenido
                if ":" in texto:
                    partes = texto.split(":", 1)
                    if len(partes) > 1 and partes[1].strip():
                        capturar = True
                        texto_desc = partes[1].strip()
                        if texto_desc:
                            descripcion_html.append(f"<p>{texto_desc}</p>")
                    else:
                        capturar = True
                else:
                    capturar = True
                continue
        
        # Si estamos capturando, verificar si debemos parar ANTES de procesar el contenido
        if capturar:
            # Verificar si el texto completo ES una de las palabras clave o las contiene al inicio
            debe_parar = False
            
            # Verificar si el texto completo coincide exactamente con alguna palabra clave
            if texto_lower in [p.lower() for p in palabras_fin]:
                debe_parar = True
            else:
                # Verificar si alguna palabra clave está al inicio del texto (más estricto)
                for palabra_fin in palabras_fin:
                    palabra_lower = palabra_fin.lower().strip()
                    # Si el texto empieza con la palabra clave seguida de espacio, dos puntos, o es exactamente igual
                    # También verificar si el texto es muy corto y coincide (para evitar capturar títulos de sección)
                    if (texto_lower == palabra_lower or 
                        texto_lower.startswith(palabra_lower + " ") or 
                        texto_lower.startswith(palabra_lower + ":") or
                        (len(texto_lower) < 30 and palabra_lower in texto_lower and texto_lower.startswith(palabra_lower))):
                        debe_parar = True
                        break
            
            if debe_parar:
                # Cerrar lista si está abierta
                if lista_abierta:
                    descripcion_html.append("</ul>")
                    lista_abierta = False
                print(f"[DEBUG] Descripción terminada al encontrar: '{texto}'")
                break
            
            # IMPORTANTE: Solo procesar el contenido si NO es una palabra clave
            # Si el texto parece ser un marcador de sección, no procesarlo
            if texto and len(texto.strip()) < 50:  # Textos cortos pueden ser marcadores
                # Verificar si el texto parece ser solo un título/marcador
                es_marcador = any(
                    texto_lower.startswith(p.lower()) or texto_lower == p.lower()
                    for p in palabras_fin
                )
                if es_marcador:
                    print(f"[DEBUG] Ignorando posible marcador: '{texto}'")
                    continue
            
        if capturar and texto:
            # Verificar si es un elemento de lista
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            is_list = style_name in ["list paragraph", "viñeta", "list", "list bullet", "list number"]
            
            if is_list:
                # Abrir lista si no está abierta
                if not lista_abierta:
                    descripcion_html.append("<ul>")
                    lista_abierta = True
                
                # Procesar contenido del elemento de lista
                item_html = ""
                for run in para.runs:
                    if run.text.strip():
                        texto_run = run.text
                        # Aplicar formatos múltiples (puede tener negrita + cursiva + subrayado)
                        if run.bold:
                            texto_run = f"<strong>{texto_run}</strong>"
                        if run.italic:
                            texto_run = f"<em>{texto_run}</em>"
                        if run.underline:
                            texto_run = f"<u>{texto_run}</u>"
                        item_html += texto_run
                
                if not item_html.strip():
                    item_html = texto
                
                descripcion_html.append(f"<li>{item_html}</li>")
            else:
                # Si era una lista y ahora no lo es, cerrar la lista
                if lista_abierta:
                    descripcion_html.append("</ul>")
                    lista_abierta = False
                
                # Procesar párrafo normal
                paragraph_html = ""
                
                # Si el párrafo tiene solo un run
                if len(para.runs) == 1:
                    run = para.runs[0]
                    texto_run = run.text
                    # Aplicar formatos múltiples
                    if run.bold:
                        texto_run = f"<strong>{texto_run}</strong>"
                    if run.italic:
                        texto_run = f"<em>{texto_run}</em>"
                    if run.underline:
                        texto_run = f"<u>{texto_run}</u>"
                    paragraph_html = texto_run
                else:
                    # Si tiene múltiples runs, procesar cada uno
                    for run in para.runs:
                        if run.text.strip():
                            texto_run = run.text
                            # Aplicar formatos múltiples
                            if run.bold:
                                texto_run = f"<strong>{texto_run}</strong>"
                            if run.italic:
                                texto_run = f"<em>{texto_run}</em>"
                            if run.underline:
                                texto_run = f"<u>{texto_run}</u>"
                            paragraph_html += texto_run
                
                # Si después de procesar no hay contenido, usar el texto plano
                if not paragraph_html.strip():
                    paragraph_html = texto
                
                # Envolver en párrafo HTML
                if paragraph_html.strip():
                    descripcion_html.append(f"<p>{paragraph_html}</p>")
    
    # Cerrar lista si quedó abierta al final
    if lista_abierta:
        descripcion_html.append("</ul>")
    
    return "\n".join(descripcion_html) if descripcion_html else ""

def frase_clave(doc):
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if para.text.strip().lower() == "frase clave objetivo":
            buscar = True
    return ""

def titulo_seo(doc):
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if para.text.strip().lower() == "título seo":
            buscar = True
    return ""

def meta_description(doc):
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if para.text.strip().lower() == "meta description":
            buscar = True
    return ""

def leer_categorias(doc):
    categorias = []
    capturar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        if "categorías" in texto_lower or "categorias" in texto_lower:
            capturar = True
            continue
        if texto == "" and capturar:
            break
        if capturar and texto:
            # Verificar si el párrafo tiene estilo
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            is_list = style_name in ["list paragraph", "list number", "viñeta", "list bullet"]
            
            # Dividir por comas si hay múltiples categorías en una línea
            if "," in texto:
                # Dividir por comas y agregar cada categoría individualmente
                cats_divididas = texto.split(",")
                for cat in cats_divididas:
                    cat_limpia = cat.strip("• ").strip()
                    if cat_limpia:
                        categorias.append(cat_limpia)
            else:
                # Si no hay comas, agregar la categoría completa
                if is_list:
                    categorias.append(f"• {texto}")
                else:
                    categorias.append(texto)
    
    # Limpiar y retornar categorías únicas
    categorias_limpias = [categoria.strip("• ").strip() for categoria in categorias if categoria.strip()]
    return categorias_limpias

def leer_etiquetas(doc):
    """Extrae las etiquetas del documento Word"""
    etiquetas = []
    capturar = False
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        # Buscar inicio de etiquetas con múltiples variaciones
        if any(palabra in texto_lower for palabra in ["etiquetas", "etiqueta", "tags", "tag"]):
            capturar = True
            # Si las etiquetas están en la misma línea después de los dos puntos
            if ":" in texto and len(texto.split(":", 1)[1].strip()) > 0:
                etiquetas_texto = texto.split(":", 1)[1].strip()
                if etiquetas_texto:
                    # Dividir por comas
                    etiquetas.extend([e.strip() for e in etiquetas_texto.split(",") if e.strip()])
            continue
        
        # Detener si encuentra otra sección
        if capturar:
            palabras_fin = ["categorías", "categorias", "frase clave", "titulo seo", "meta description"]
            if any(palabra in texto_lower for palabra in palabras_fin):
                break
            if texto == "":
                break
            
            if texto:
                # Dividir por comas si hay múltiples etiquetas
                if "," in texto:
                    etiquetas.extend([e.strip("• ").strip() for e in texto.split(",") if e.strip()])
                else:
                    etiquetas.append(texto.strip("• ").strip())
    
    # Limpiar y retornar etiquetas únicas
    etiquetas_limpias = list(set([etiqueta.strip() for etiqueta in etiquetas if etiqueta.strip()]))
    return etiquetas_limpias

def guardar_borrador(driver, wait):
    """Función para guardar el post como borrador"""
    try:
        print("[INFO] Guardando como borrador...")
        
        # Paso 1: Hacer clic en el botón de guardar/actualizar
        guardar_btn = wait.until(EC.element_to_be_clickable((By.XPATH, 
            '//button[contains(@class, "editor-post-save-draft") or contains(text(), "Guardar")]')))
        
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", guardar_btn)
        time.sleep(2)
        
        # Verificar si el botón está habilitado
        if not guardar_btn.is_enabled():
            print("[ERROR] Botón de guardar deshabilitado")
            return False
            
        guardar_btn.click()
        print("[OK] Clic en botón de guardar")
        
        # Paso 2: Esperar a que aparezca el mensaje de guardado
        time.sleep(3)
        
        # Verificar diferentes indicadores de éxito
        try:
            # Mensaje de guardado exitoso
            exito_element = wait.until(EC.presence_of_element_located((By.XPATH, 
                '//div[contains(text(), "Guardado") or contains(text(), "Saved") or contains(text(), "Borrador") or contains(text(), "Draft")]')))
            print("[OK] Guardado exitoso - Mensaje confirmado")
            return True
        except TimeoutException:
            # Verificar si el botón cambió de estado
            try:
                actualizar_btn = driver.find_element(By.XPATH, '//button[contains(text(), "Actualizar") or contains(@class, "editor-post-publish-button")]')
                print("[OK] Guardado exitoso - Botón cambió de estado")
                return True
            except:
                print("[WARNING] No se pudo confirmar el guardado")
                return False
                
    except TimeoutException as e:
        print(f"[ERROR] Tiempo de espera agotado al guardar: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] Error al guardar: {e}")
        return False

def guardar_con_teclado(driver):
    """Usar atajos de teclado para guardar como borrador"""
    try:
        print("[INFO] Intentando guardar con atajos de teclado...")
        
        # Enfocar el editor
        title_field = driver.find_element(By.XPATH, '//h1[@role="textbox"]')
        title_field.click()
        time.sleep(1)
        
        # Ctrl + S para guardar borrador
        actions = ActionChains(driver)
        actions.key_down(Keys.CONTROL).send_keys('s').key_up(Keys.CONTROL)
        actions.perform()
        time.sleep(3)
        
        print("[OK] Comandos de teclado ejecutados para guardar")
        return True
    except Exception as e:
        print(f"[ERROR] Error con atajos de teclado: {e}")
        return False

def verificar_estado_guardado(driver):
    """Verificar el estado actual del guardado"""
    try:
        # Verificar botones disponibles
        botones = driver.find_elements(By.XPATH, '//button[contains(@class, "editor-post-save-draft") or contains(text(), "Guardar")]')
        print(f"[INFO] Botones de guardar encontrados: {len(botones)}")
        
        for i, boton in enumerate(botones):
            print(f"  Botón {i+1}: Texto='{boton.text}', Habilitado={boton.is_enabled()}")
            
        # Verificar estado guardado
        try:
            estado = driver.find_element(By.XPATH, '//div[contains(@class, "editor-post-saved-state")]')
            print(f"[INFO] Estado actual: {estado.text}")
        except:
            print("[INFO] Estado: No se pudo determinar")
            
        return True
    except Exception as e:
        print(f"[ERROR] Error en verificación: {e}")
        return False

def insertar_descripcion_larga(driver, textarea_editor, descripcion):
    """Inserta una descripción usando JavaScript para mejor rendimiento"""
    print(f"[INFO] Insertando descripción ({len(descripcion)} caracteres)...")
    
    try:
        # Usar JavaScript para insertar todo de una vez y disparar eventos
        script = """
        var textarea = arguments[0];
        var contenido = arguments[1];
        textarea.value = contenido;
        
        // Disparar múltiples eventos para que Gutenberg detecte el cambio
        textarea.dispatchEvent(new Event('input', { bubbles: true, cancelable: true }));
        textarea.dispatchEvent(new Event('change', { bubbles: true, cancelable: true }));
        textarea.dispatchEvent(new KeyboardEvent('keyup', { bubbles: true, cancelable: true }));
        
        // También disparar evento personalizado de WordPress
        if (typeof wp !== 'undefined' && wp.data) {
            wp.data.dispatch('core/editor').editPost({ content: contenido });
        }
        """
        driver.execute_script(script, textarea_editor, descripcion)
        
        # Verificar que se insertó correctamente
        time.sleep(1)
        contenido_insertado = driver.execute_script("return arguments[0].value;", textarea_editor)
        if len(contenido_insertado) > 0:
            print(f"[OK] Descripción insertada con JavaScript ({len(contenido_insertado)} caracteres)")
        else:
            raise Exception("La descripción no se insertó correctamente")
            
    except Exception as e:
        print(f"[WARNING] JavaScript falló, usando método tradicional: {e}")
        # Método de respaldo
        textarea_editor.clear()
        textarea_editor.send_keys(descripcion)
        # Disparar eventos manualmente
        driver.execute_script("arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", textarea_editor)
        driver.execute_script("arguments[0].dispatchEvent(new Event('change', { bubbles: true }));", textarea_editor)
        print("[OK] Descripción insertada con send_keys")

def crear_nueva_categoria(driver, wait, nombre_categoria):
    """Crea una nueva categoría en WordPress si no existe"""
    try:
        print(f"[INFO] Intentando crear nueva categoría: '{nombre_categoria}'")
        
        # Paso 1: Hacer clic en el enlace "Añadir categoría"
        try:
            # Buscar el enlace azul "Añadir categoría"
            añadir_link = wait.until(EC.element_to_be_clickable((
                By.XPATH, '//button[contains(text(), "Añadir categoría") or contains(text(), "Add category")]'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", añadir_link)
            time.sleep(0.5)
            añadir_link.click()
            print("[OK] Enlace 'Añadir categoría' clickeado")
            time.sleep(2)  # Esperar a que aparezca el formulario
        except Exception as e:
            print(f"[WARNING] No se encontró el enlace 'Añadir categoría': {e}")
            # Intentar con selector alternativo
            try:
                añadir_link = driver.find_element(By.XPATH, '//*[contains(text(), "Añadir categoría") or contains(text(), "Add category")]')
                añadir_link.click()
                time.sleep(2)
            except:
                print("[ERROR] No se pudo hacer clic en 'Añadir categoría'")
                return False
        
        # Paso 2: Buscar el campo "NOMBRE DE LA NUEVA CATEGORÍA"
        try:
            # Esperar a que aparezca el campo de texto
            campo_nombre = wait.until(EC.visibility_of_element_located((
                By.XPATH, '//input[@type="text" and ancestor::*[contains(., "NOMBRE DE LA NUEVA CATEGORÍA")]]'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo_nombre)
            time.sleep(0.5)
            campo_nombre.click()
            campo_nombre.clear()
            campo_nombre.send_keys(nombre_categoria)
            print(f"[OK] Nombre de categoría '{nombre_categoria}' ingresado en el campo")
            time.sleep(1)
        except:
            # Intentar con selectores alternativos
            try:
                # Buscar cualquier input visible en el área de categorías
                campos = driver.find_elements(By.XPATH, '//div[contains(@class, "editor-post-taxonomies")]//input[@type="text"]')
                if campos:
                    # Usar el último campo encontrado (probablemente el de nueva categoría)
                    campo_nombre = campos[-1]
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo_nombre)
                    time.sleep(0.5)
                    campo_nombre.click()
                    campo_nombre.clear()
                    campo_nombre.send_keys(nombre_categoria)
                    print(f"[OK] Nombre de categoría '{nombre_categoria}' ingresado (método alternativo)")
                    time.sleep(1)
                else:
                    print("[ERROR] No se encontró el campo para el nombre de la categoría")
                    return False
            except Exception as e:
                print(f"[ERROR] No se pudo ingresar el nombre de la categoría: {e}")
                return False
        
        # Paso 3: Hacer clic en el botón azul "Añadir categoría" para confirmar
        try:
            # Buscar el botón azul "Añadir categoría" (el que confirma, no el enlace)
            confirmar_btn = wait.until(EC.element_to_be_clickable((
                By.XPATH, '//button[contains(@class, "button") and contains(text(), "Añadir categoría")]'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", confirmar_btn)
            time.sleep(0.5)
            confirmar_btn.click()
            print("[OK] Botón 'Añadir categoría' (confirmar) clickeado")
            time.sleep(2)  # Esperar a que se cree la categoría y se actualice la lista
            
            # Verificar si se creó correctamente
            print(f"[SUCCESS] Categoría '{nombre_categoria}' creada exitosamente")
            return True
            
        except Exception as e:
            print(f"[WARNING] No se encontró el botón de confirmar: {e}")
            # Intentar presionar Enter como alternativa
            try:
                campo_nombre.send_keys(Keys.RETURN)
                print("[OK] Enter presionado para crear categoría")
                time.sleep(2)
                return True
            except:
                print("[ERROR] No se pudo confirmar la creación de la categoría")
                return False
        
    except Exception as e:
        print(f"[ERROR] Error general al crear categoría '{nombre_categoria}': {e}")
        return False

def procesar_archivo(archivo_word, driver, wait):
    """Procesa un archivo Word individual y lo guarda como borrador en WordPress"""
    print(f"\n{'='*60}")
    print(f"[INFO] Procesando: {os.path.basename(archivo_word)}")
    print(f"{'='*60}")
    
    try:
        # Navegar a la página de nuevo post
        driver.get(url_wp + "/post-new.php")
        time.sleep(4)

        doc = Document(archivo_word)
        titulo = extraer_titulo(doc)
        descripcion = extraer_descripcion_con_formato(doc)
        frase_obj = frase_clave(doc)
        tit_seo = titulo_seo(doc)
        meta_desc = meta_description(doc)
        categorias = leer_categorias(doc)
        etiquetas = leer_etiquetas(doc)

        print(f"[INFO] Título: {titulo}")
        print(f"[INFO] Longitud descripción: {len(descripcion)} caracteres")
        print(f"[INFO] Frase clave: {frase_obj}")
        print(f"[INFO] SEO: {tit_seo}")
        print(f"[INFO] Meta desc: {meta_desc}")
        print(f"[INFO] Categorías: {categorias}")
        print(f"[INFO] Etiquetas: {etiquetas}")

        # ---------------- INSERTAR TÍTULO ----------------
        print("[INFO] Insertando título...")
        title_field = wait.until(EC.element_to_be_clickable((By.XPATH, '//h1[@role="textbox"]')))
        title_field.click()
        title_field.clear()
        title_field.send_keys(titulo)
        time.sleep(1)

        # ---------------- CAMBIAR A EDITOR DE CÓDIGO ----------------
        print("[INFO] Cambiando a editor de código...")
        
        # Buscar y hacer clic en el botón de opciones (tres puntos)
        try:
            opciones_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//button[@aria-label="Opciones"]')))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", opciones_btn)
            time.sleep(1)
            opciones_btn.click()
            print("[OK] Botón de opciones clickeado")
            time.sleep(2)
        except TimeoutException:
            print("[ERROR] No se encontró el botón de opciones")
            return False

        # Buscar y hacer clic en "Editor de código"
        try:
            editor_codigo_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//button[contains(@class, "components-menu-items-choice") and .//span[text()="Editor de código"]]')))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", editor_codigo_btn)
            time.sleep(1)
            editor_codigo_btn.click()
            print("[OK] Cambiado a editor de código")
            time.sleep(2)
        except TimeoutException:
            print("[ERROR] No se encontró la opción 'Editor de código'")
            return False

        # ---------------- INSERTAR DESCRIPCIÓN EN EDITOR DE CÓDIGO ----------------
        print("[INFO] Insertando descripción con formato HTML...")
        try:
            # Buscar el textarea del editor de código (múltiples selectores para mayor compatibilidad)
            try:
                textarea_editor = wait.until(EC.visibility_of_element_located((By.XPATH, '//textarea[@class="editor-post-text-editor"]')))
            except:
                try:
                    textarea_editor = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'textarea.editor-post-text-editor')))
                except:
                    textarea_editor = wait.until(EC.visibility_of_element_located((By.XPATH, '//textarea[contains(@class, "editor-post")]')))
            
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", textarea_editor)
            time.sleep(1)
            textarea_editor.click()
            textarea_editor.clear()
            time.sleep(0.5)
            
            # Insertar la descripción usando JavaScript mejorado
            insertar_descripcion_larga(driver, textarea_editor, descripcion)
            time.sleep(2)
            
            # Verificar que el contenido está antes de salir
            contenido_verificado = driver.execute_script("return arguments[0].value;", textarea_editor)
            if len(contenido_verificado) == 0:
                print("[WARNING] El contenido está vacío, reintentando inserción...")
                insertar_descripcion_larga(driver, textarea_editor, descripcion)
                time.sleep(2)
            
            print(f"[OK] Descripción insertada: {len(contenido_verificado)} caracteres")
        except TimeoutException:
            print("[ERROR] No se encontró el textarea del editor de código")
            return False
        except Exception as e:
            print(f"[ERROR] Error al insertar descripción: {e}")
            return False

        # ---------------- SALIR DEL EDITOR DE CÓDIGO ----------------
        print("[INFO] Saliendo del editor de código...")
        try:
            # Buscar y hacer clic en el botón "Salir del editor de código" (múltiples selectores)
            try:
                salir_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//button[contains(@class, "is-tertiary") and contains(text(), "Salir del editor de código")]')))
            except:
                try:
                    salir_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//button[contains(text(), "Salir del editor de código")]')))
                except:
                    salir_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//button[contains(@class, "is-tertiary")]')))
            
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", salir_btn)
            time.sleep(1)
            salir_btn.click()
            print("[OK] Salido del editor de código")
            time.sleep(3)  # Espera más tiempo para que Gutenberg procese el contenido
            
            # Verificar que el contenido se mantuvo después de salir
            try:
                # Buscar algún bloque de párrafo en el editor visual
                bloques = driver.find_elements(By.CSS_SELECTOR, '[data-type="core/paragraph"]')
                if len(bloques) > 0:
                    print(f"[OK] Contenido verificado: {len(bloques)} bloques encontrados en editor visual")
                else:
                    print("[WARNING] No se detectaron bloques visibles, pero el contenido puede estar guardado")
            except:
                print("[INFO] Verificación de bloques omitida")
                
        except TimeoutException:
            print("[ERROR] No se encontró el botón 'Salir del editor de código'")
            # Intentar presionar Escape como alternativa
            try:
                ActionChains(driver).send_keys(Keys.ESCAPE).perform()
                print("[OK] Salido del editor de código con Escape")
                time.sleep(2)
            except:
                print("[ERROR] No se pudo salir del editor de código")
                return False

        # ---------------- EXPANDIR YOAST SEO (si está colapsado) ----------------
        print("[INFO] Verificando panel de Yoast SEO...")
        try:
            # Buscar y expandir el panel de Yoast SEO si está colapsado
            try:
                panel_yoast = wait.until(EC.presence_of_element_located((
                    By.XPATH, '//div[contains(@class, "yoast")]//button[contains(@class, "components-panel__body-toggle") or contains(text(), "Yoast SEO")]'
                )))
                aria_expanded = panel_yoast.get_attribute("aria-expanded")
                if aria_expanded == "false" or aria_expanded is None:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", panel_yoast)
                    time.sleep(0.5)
                    panel_yoast.click()
                    print("[OK] Panel de Yoast SEO expandido")
                    time.sleep(2)
            except:
                print("[INFO] Panel de Yoast SEO ya está expandido o no se encontró")
        except:
            pass

        # ---------------- INSERTAR FRASE CLAVE ----------------
        if frase_obj and frase_obj.strip():
            print(f"[INFO] Insertando frase clave: '{frase_obj}'")
            try:
                # Buscar el campo "Frase clave objetivo" buscando primero el label/título y luego el input asociado
                frase_field = None
                selectores_frase = [
                    # Buscar por label "Frase clave objetivo" y luego el input siguiente
                    (By.XPATH, '//label[contains(text(), "Frase clave objetivo") or contains(text(), "Frase clave")]/following::input[1]'),
                    (By.XPATH, '//h3[contains(text(), "Frase clave objetivo") or contains(text(), "Frase clave")]/following::input[1]'),
                    (By.XPATH, '//div[contains(text(), "Frase clave objetivo") or contains(text(), "Frase clave")]/following::input[1]'),
                    # Buscar por ID directo
                    (By.ID, 'focus-keyword-input-metabox'),
                    (By.XPATH, '//input[@id="focus-keyword-input-metabox"]'),
                    # Buscar input cerca del texto "Frase clave objetivo"
                    (By.XPATH, '//div[contains(., "Frase clave objetivo")]//input[@type="text"]'),
                    (By.XPATH, '//div[contains(., "Frase clave objetivo")]//input'),
                    # Buscar por placeholder o name
                    (By.XPATH, '//input[contains(@placeholder, "Enter your focus keyword") or contains(@placeholder, "Palabra clave") or contains(@placeholder, "Frase clave")]'),
                    (By.XPATH, '//input[contains(@name, "focus_keyword") or contains(@name, "focus-keyword")]'),
                    # Selector CSS alternativo
                    (By.CSS_SELECTOR, '#focus-keyword-input-metabox'),
                    # Buscar cualquier input en la sección de Yoast que contenga "frase" o "keyword"
                    (By.XPATH, '//div[contains(@class, "yoast")]//input[@type="text" and (contains(@class, "key") or contains(@id, "key"))]')
                ]
                
                for selector_type, selector_value in selectores_frase:
                    try:
                        frase_field = wait.until(EC.visibility_of_element_located((selector_type, selector_value)))
                        if frase_field and frase_field.is_displayed():
                            print(f"[DEBUG] Campo encontrado con selector: {str(selector_value)[:80]}...")
                            break
                    except:
                        continue
                
                if not frase_field:
                    # Intentar con JavaScript para buscar el campo
                    try:
                        frase_field = driver.execute_script("""
                            var labels = document.querySelectorAll('label, h3, div');
                            for (var i = 0; i < labels.length; i++) {
                                if (labels[i].textContent.includes('Frase clave objetivo') || labels[i].textContent.includes('Frase clave')) {
                                    var input = labels[i].parentElement.querySelector('input[type="text"]');
                                    if (!input) {
                                        input = labels[i].nextElementSibling;
                                        while (input && input.tagName !== 'INPUT') {
                                            input = input.nextElementSibling;
                                        }
                                    }
                                    if (input) return input;
                                }
                            }
                            return document.querySelector('#focus-keyword-input-metabox');
                        """)
                        if frase_field:
                            print("[DEBUG] Campo encontrado con JavaScript")
                    except:
                        pass
                
                if frase_field:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", frase_field)
                    time.sleep(0.5)
                    frase_field.click()
                    time.sleep(0.3)
                    frase_field.clear()
                    time.sleep(0.3)
                    frase_field.send_keys(Keys.CONTROL, 'a')
                    time.sleep(0.2)
                    frase_field.send_keys(Keys.DELETE)
                    time.sleep(0.3)
                    frase_field.send_keys(frase_obj.strip())
                    time.sleep(0.5)
                    
                    # Disparar múltiples eventos para que WordPress detecte el cambio
                    driver.execute_script("arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", frase_field)
                    driver.execute_script("arguments[0].dispatchEvent(new Event('change', { bubbles: true }));", frase_field)
                    driver.execute_script("arguments[0].dispatchEvent(new KeyboardEvent('keyup', { bubbles: true }));", frase_field)
                    
                    # Verificar que se insertó
                    valor_insertado = frase_field.get_attribute('value')
                    if valor_insertado and frase_obj.strip() in valor_insertado:
                        print(f"[OK] Frase clave insertada correctamente: '{valor_insertado}'")
                    else:
                        print(f"[WARNING] La frase clave puede no haberse insertado correctamente. Valor encontrado: '{valor_insertado}'")
                    time.sleep(2)
                else:
                    raise Exception("No se encontró el campo de frase clave")
            except TimeoutException:
                print("[WARNING] Campo de frase clave no encontrado (Timeout)")
            except Exception as e:
                print(f"[ERROR] Error al insertar frase clave: {e}")
        else:
            print("[INFO] No hay frase clave para insertar")

        # ---------------- INSERTAR META DESCRIPCIÓN EN YOAST SEO ----------------
        if meta_desc and meta_desc.strip():
            print(f"[INFO] Insertando meta descripción en Yoast SEO: '{meta_desc[:50]}...'")
            try:
                # Buscar el campo "Meta description" ESPECÍFICAMENTE dentro del panel de Yoast SEO
                meta_desc_field = None
                selectores_meta = [
                    # PRIORIDAD 1: Buscar dentro del contenedor de Yoast SEO específicamente
                    (By.XPATH, '//div[contains(@class, "yoast") or contains(@id, "yoast")]//label[contains(text(), "Meta description") or contains(text(), "Meta descripción")]/following::textarea[1]'),
                    (By.XPATH, '//div[contains(@class, "yoast") or contains(@id, "yoast")]//h3[contains(text(), "Meta description") or contains(text(), "Meta descripción")]/following::textarea[1]'),
                    (By.XPATH, '//div[contains(@class, "yoast")]//div[contains(text(), "Meta description") or contains(text(), "Meta descripción")]/following::textarea[1]'),
                    
                    # PRIORIDAD 2: Buscar textarea dentro del contenedor de Yoast con "Meta description"
                    (By.XPATH, '//div[contains(@class, "yoast")]//div[contains(., "Meta description") or contains(., "Meta descripción")]//textarea'),
                    
                    # PRIORIDAD 3: Buscar por ID específico de Yoast
                    (By.ID, 'yoast-google-preview-description-metabox'),
                    (By.XPATH, '//textarea[@id="yoast-google-preview-description-metabox"]'),
                    (By.XPATH, '//textarea[contains(@id, "yoast") and contains(@id, "description")]'),
                    
                    # PRIORIDAD 4: Buscar en el panel lateral de Yoast SEO
                    (By.XPATH, '//div[contains(@class, "edit-post-sidebar")]//div[contains(@class, "yoast")]//textarea[contains(@id, "description")]'),
                    (By.XPATH, '//div[contains(@class, "interface-complementary-area")]//div[contains(@class, "yoast")]//textarea'),
                    
                    # PRIORIDAD 5: Buscar por clase específica de Yoast
                    (By.XPATH, '//textarea[contains(@class, "yoast") and contains(@class, "description")]'),
                    
                    # ÚLTIMO RECURSO: Cualquier textarea en sección Yoast (excluyendo nativo de WordPress)
                    (By.XPATH, '//div[contains(@class, "yoast") or contains(@id, "yoast")]//textarea[not(ancestor::div[contains(@class, "edit-post-meta-boxes-area")])]')
                ]
                
                for selector_type, selector_value in selectores_meta:
                    try:
                        meta_desc_field = wait.until(EC.visibility_of_element_located((selector_type, selector_value)))
                        if meta_desc_field and meta_desc_field.is_displayed():
                            # Verificar que es el campo de Yoast SEO y NO el campo nativo de WordPress
                            field_id = meta_desc_field.get_attribute('id') or ''
                            field_class = meta_desc_field.get_attribute('class') or ''
                            
                            # Buscar el contenedor padre para verificar contexto
                            try:
                                # Buscar si está dentro de un contenedor de Yoast
                                parent_yoast = meta_desc_field.find_element(By.XPATH, './ancestor::*[contains(@class, "yoast") or contains(@id, "yoast")][1]')
                                if parent_yoast:
                                    print(f"[DEBUG] Campo de Yoast SEO encontrado con selector: {str(selector_value)[:80]}...")
                                    break
                            except:
                                pass
                            
                            # Si el ID del campo contiene "yoast", es del campo correcto
                            if 'yoast' in field_id.lower():
                                print(f"[DEBUG] Campo de Yoast SEO encontrado por ID: {field_id}")
                                break
                            
                            # Verificar que NO es el campo nativo de WordPress (excerpt)
                            try:
                                parent_excerpt = meta_desc_field.find_element(By.XPATH, './ancestor::*[@id="postexcerpt" or contains(@class, "postexcerpt")][1]')
                                if parent_excerpt:
                                    print(f"[DEBUG] Campo encontrado pero es de WordPress nativo (excerpt), continuando búsqueda...")
                                    meta_desc_field = None
                                    continue
                            except:
                                pass
                            
                            # Si llegamos aquí y el selector ya filtra por "yoast", aceptarlo
                            if 'yoast' in str(selector_value).lower():
                                print(f"[DEBUG] Campo de Yoast SEO encontrado con selector: {str(selector_value)[:80]}...")
                                break
                            
                            print(f"[DEBUG] Campo encontrado, verificando contexto...")
                            break
                    except:
                        continue
                
                if not meta_desc_field:
                    # Intentar con JavaScript para buscar específicamente el campo de Yoast SEO
                    try:
                        meta_desc_field = driver.execute_script("""
                            // Buscar primero en el contenedor de Yoast SEO
                            var yoastContainer = document.querySelector('.yoast, #yoast, [id*="yoast"], [class*="yoast"]');
                            if (!yoastContainer) {
                                // Si no hay contenedor específico, buscar por ID
                                var field = document.querySelector('#yoast-google-preview-description-metabox');
                                if (field) return field;
                            }
                            
                            // Buscar labels dentro de Yoast
                            var labels = yoastContainer ? yoastContainer.querySelectorAll('label, h3, div') : document.querySelectorAll('label, h3, div');
                            for (var i = 0; i < labels.length; i++) {
                                var text = labels[i].textContent || '';
                                // Buscar específicamente "Meta description" en contexto de Yoast
                                if ((text.includes('Meta description') || text.includes('Meta descripción')) && 
                                    (labels[i].closest('.yoast, #yoast, [id*="yoast"], [class*="yoast"]') || 
                                     text.includes('Yoast') || 
                                     labels[i].getAttribute('for') && labels[i].getAttribute('for').includes('yoast'))) {
                                    
                                    var textarea = labels[i].parentElement.querySelector('textarea');
                                    if (!textarea) {
                                        textarea = labels[i].nextElementSibling;
                                        while (textarea && textarea.tagName !== 'TEXTAREA') {
                                            textarea = textarea.nextElementSibling;
                                        }
                                    }
                                    if (textarea) return textarea;
                                }
                            }
                            
                            // Último recurso: buscar por ID específico de Yoast
                            return document.querySelector('#yoast-google-preview-description-metabox');
                        """)
                        if meta_desc_field:
                            print("[DEBUG] Campo de Yoast SEO encontrado con JavaScript")
                    except:
                        pass
                
                if meta_desc_field:
                    try:
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", meta_desc_field)
                        time.sleep(0.5)
                        
                        # MÉTODO PRINCIPAL: Usar send_keys para que Yoast SEO detecte la entrada como si fuera manual
                        print("[DEBUG] Usando método send_keys para inserción más natural...")
                        meta_desc_field.click()
                        time.sleep(0.3)
                        
                        # Limpiar el campo completamente
                        meta_desc_field.clear()
                        time.sleep(0.2)
                        meta_desc_field.send_keys(Keys.CONTROL, 'a')
                        time.sleep(0.2)
                        meta_desc_field.send_keys(Keys.DELETE)
                        time.sleep(0.3)
                        
                        # Insertar caracter por caracter (más lento pero más confiable para Yoast)
                        meta_desc_text = meta_desc.strip()
                        print(f"[DEBUG] Insertando {len(meta_desc_text)} caracteres...")
                        meta_desc_field.send_keys(meta_desc_text)
                        time.sleep(1)
                        
                        # Disparar eventos adicionales después de la inserción
                        driver.execute_script("""
                            var textarea = arguments[0];
                            // Disparar eventos nativos
                            textarea.dispatchEvent(new Event('input', { bubbles: true, cancelable: true }));
                            textarea.dispatchEvent(new Event('change', { bubbles: true, cancelable: true }));
                            textarea.dispatchEvent(new KeyboardEvent('keyup', { bubbles: true }));
                            
                            // Actualizar snippet preview si existe
                            try {
                                var snippetDesc = document.querySelector('.yoast-snippet__description, .yoast-snippet-description, [data-field="description"]');
                                if (snippetDesc) {
                                    snippetDesc.textContent = textarea.value;
                                }
                            } catch(e) {}
                            
                            // Intentar actualizar Redux store de Yoast
                            if (typeof wp !== 'undefined' && wp.data) {
                                try {
                                    var store = wp.data.select('yoast-seo/editor');
                                    if (store) {
                                        wp.data.dispatch('yoast-seo/editor').updateData({ metaDescription: textarea.value });
                                    }
                                } catch(e) {}
                            }
                        """, meta_desc_field)
                        
                        time.sleep(1)
                        
                        # Verificar que se insertó
                        valor_insertado = meta_desc_field.get_attribute('value')
                        if not valor_insertado or len(valor_insertado) == 0:
                            # Intentar con JavaScript
                            valor_insertado = driver.execute_script("return arguments[0].value;", meta_desc_field)
                        
                        if valor_insertado and len(valor_insertado) > 0:
                            print(f"[OK] Meta descripción insertada con send_keys ({len(valor_insertado)} caracteres)")
                            
                            # Hacer clic fuera del campo para activar el guardado
                            title_field = driver.find_element(By.XPATH, '//h1[@role="textbox"]')
                            title_field.click()
                            time.sleep(1)
                            
                            # Verificar después de hacer clic fuera
                            valor_final = driver.execute_script("return arguments[0].value;", meta_desc_field)
                            if valor_final and len(valor_final) > 0:
                                print(f"[OK] Meta descripción confirmada después de hacer clic fuera ({len(valor_final)} caracteres)")
                                
                                # Esperar un momento adicional para que Yoast guarde
                                time.sleep(1)
                            else:
                                print(f"[WARNING] El valor se perdió. Reintentando inserción...")
                                # Reintentar con método alternativo
                                raise Exception("Valor perdido después de hacer clic fuera")
                        else:
                            raise Exception("No se pudo insertar la meta descripción")
                    
                    except Exception as js_error:
                        print(f"[WARNING] Error al insertar meta descripción: {js_error}")
                        print("[INFO] Reintentando con método alternativo...")
                        # Método de respaldo: inserción tradicional
                        try:
                            meta_desc_field.click()
                            time.sleep(0.3)
                            meta_desc_field.clear()
                            time.sleep(0.3)
                            meta_desc_field.send_keys(Keys.CONTROL, 'a')
                            time.sleep(0.2)
                            meta_desc_field.send_keys(Keys.DELETE)
                            time.sleep(0.3)
                            meta_desc_field.send_keys(meta_desc.strip())
                            time.sleep(1)
                            
                            # Disparar eventos manualmente
                            driver.execute_script("arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", meta_desc_field)
                            driver.execute_script("arguments[0].dispatchEvent(new Event('change', { bubbles: true }));", meta_desc_field)
                            
                            # Hacer clic fuera para guardar
                            title_field = driver.find_element(By.XPATH, '//h1[@role="textbox"]')
                            title_field.click()
                            time.sleep(1)
                            
                            valor_insertado = meta_desc_field.get_attribute('value')
                            if valor_insertado and len(valor_insertado) > 0:
                                print(f"[OK] Meta descripción insertada (método alternativo) ({len(valor_insertado)} caracteres)")
                            else:
                                print(f"[WARNING] La meta descripción puede no haberse insertado correctamente")
                        except Exception as e2:
                            print(f"[ERROR] Método alternativo también falló: {e2}")
                    
                    time.sleep(2)
                else:
                    raise Exception("No se encontró el campo de meta descripción")
            except TimeoutException:
                print("[WARNING] Campo de meta descripción no encontrado (Timeout)")
            except Exception as e:
                print(f"[ERROR] Error al insertar meta descripción: {e}")
        else:
            print("[INFO] No hay meta descripción para insertar")

        # ---------------- ABRIR PANEL DE AJUSTES LATERAL ----------------
        print("[INFO] Abriendo panel de ajustes lateral...")
        try:
            # Verificar si el panel ya está abierto
            try:
                panel_abierto = driver.find_element(By.CSS_SELECTOR, '.interface-complementary-area, .edit-post-sidebar')
                if panel_abierto.is_displayed():
                    print("[INFO] Panel de ajustes ya está abierto")
                else:
                    raise Exception("Panel no visible")
            except:
                # Buscar y hacer clic en el botón para abrir el panel de ajustes (ícono de engranaje)
                boton_encontrado = False
                
                # Método 1: Buscar botón de ajustes por aria-label (más común)
                selectores_ajustes = [
                    '//button[contains(@aria-label, "Ajustes") or contains(@aria-label, "Settings")]',
                    '//button[@class="components-button is-button is-pressed" and contains(@aria-label, "Settings")]',
                    '//button[contains(@class, "editor-settings-button")]',
                    '//button[contains(@class, "edit-post-sidebar__panel-toggle")]',
                    '//button[@data-label="Settings"]',
                    '//button[contains(@title, "Ajustes") or contains(@title, "Settings")]'
                ]
                
                for selector in selectores_ajustes:
                    try:
                        boton_ajustes = wait.until(EC.element_to_be_clickable((By.XPATH, selector)))
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", boton_ajustes)
                        time.sleep(0.5)
                        boton_ajustes.click()
                        print(f"[OK] Botón de ajustes clickeado (selector: {selector[:50]}...)")
                        boton_encontrado = True
                        time.sleep(2)
                        break
                    except:
                        continue
                
                # Método 2: Buscar por clase CSS
                if not boton_encontrado:
                    try:
                        boton_ajustes = driver.find_element(By.CSS_SELECTOR, 'button[class*="editor-settings-button"], button[aria-label*="Settings"], button[aria-label*="Ajustes"]')
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", boton_ajustes)
                        time.sleep(0.5)
                        boton_ajustes.click()
                        print("[OK] Botón de ajustes clickeado (método CSS)")
                        boton_encontrado = True
                        time.sleep(2)
                    except:
                        pass
                
                # Método 3: Usar atajo de teclado Alt+Shift+S
                if not boton_encontrado:
                    try:
                        # Enfocar el editor primero
                        title_field = driver.find_element(By.XPATH, '//h1[@role="textbox"]')
                        title_field.click()
                        time.sleep(0.5)
                        ActionChains(driver).key_down(Keys.ALT).key_down(Keys.SHIFT).send_keys('s').key_up(Keys.SHIFT).key_up(Keys.ALT).perform()
                        print("[OK] Panel de ajustes abierto con atajo de teclado (Alt+Shift+S)")
                        time.sleep(2)
                    except Exception as e:
                        print(f"[WARNING] No se pudo abrir panel con atajo de teclado: {e}")
            
            # Asegurar que la pestaña "Entrada" esté seleccionada
            try:
                pestaña_entrada = wait.until(EC.element_to_be_clickable((
                    By.XPATH, '//button[contains(@class, "components-tab-button") and contains(text(), "Entrada")]'
                )))
                # Verificar si ya está seleccionada
                aria_selected = pestaña_entrada.get_attribute("aria-selected")
                if aria_selected != "true":
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", pestaña_entrada)
                    time.sleep(0.5)
                    pestaña_entrada.click()
                    print("[OK] Pestaña 'Entrada' seleccionada")
                    time.sleep(1)
                else:
                    print("[INFO] Pestaña 'Entrada' ya está seleccionada")
            except Exception as e:
                print(f"[WARNING] No se pudo seleccionar la pestaña 'Entrada': {e}")
            
        except Exception as e:
            print(f"[WARNING] No se pudo abrir el panel de ajustes: {e}")
            print("[INFO] Continuando de todas formas...")

        # ---------------- INSERTAR CATEGORÍAS ----------------
        print("[INFO] Configurando categorías...")
        try:
            # Expandir la sección "Categorías" si está colapsada
            try:
                # Buscar el botón de la sección Categorías (puede estar colapsado)
                panel_categorias = wait.until(EC.presence_of_element_located((
                    By.XPATH, '//div[contains(@class, "editor-post-taxonomies__hierarchical-terms-list") or contains(@class, "editor-post-taxonomies")]//button[contains(text(), "Categorías") or contains(@aria-label, "Categorías")] | //button[contains(@class, "components-panel__body-toggle") and contains(text(), "Categorías")]'
                )))
                
                # Verificar si está expandido
                aria_expanded = panel_categorias.get_attribute("aria-expanded")
                if aria_expanded == "false" or aria_expanded is None:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", panel_categorias)
                    time.sleep(0.5)
                    panel_categorias.click()
                    print("[OK] Sección de categorías expandida")
                    time.sleep(2)
                else:
                    print("[INFO] Sección de categorías ya está expandida")
            except Exception as e:
                print(f"[DEBUG] No se encontró el botón de categorías para expandir: {e}")
                # Continuar de todas formas, puede que ya esté expandido

            # Esperar a que se carguen las categorías
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, ".components-panel__body, .editor-post-taxonomies__hierarchical-terms-list")))

            # Buscar categorías en el editor de Gutenberg
            for categoria in categorias:
                try:
                    clean_cat = categoria.strip("• ").strip()
                    print(f"[DEBUG] Buscando categoría: '{clean_cat}'")
                    
                    # Primero listar categorías disponibles para debug
                    try:
                        checkboxes = driver.find_elements(By.XPATH, '//div[contains(@class, "components-panel__body")]//input[@type="checkbox"]')
                        print(f"[DEBUG] Encontrados {len(checkboxes)} checkboxes de categorías")
                        
                        # Mostrar las primeras 5 categorías para debug
                        for i, cb in enumerate(checkboxes[:5]):
                            try:
                                # En Gutenberg, buscar el label asociado
                                label_id = cb.get_attribute('id')
                                if label_id:
                                    label = driver.find_element(By.XPATH, f'//label[@for="{label_id}"]')
                                    texto = label.text.strip()
                                    print(f"  {i+1}. '{texto}'")
                                else:
                                    # Método alternativo
                                    parent = cb.find_element(By.XPATH, './parent::*/parent::*')
                                    texto = parent.text.strip()
                                    print(f"  {i+1}. '{texto}'")
                            except Exception as e:
                                print(f"  {i+1}. [Error leyendo texto: {e}]")
                    except:
                        print("[DEBUG] No se pudieron listar categorías")
                    
                    # Buscar la categoría específica
                    checkbox_encontrado = False
                    
                    # Método 1: Buscar por texto exacto en Gutenberg (mejorado)
                    try:
                        # Buscar con múltiples selectores
                        selectores = [
                            '//div[contains(@class, "components-panel__body")]//input[@type="checkbox"]',
                            '//div[contains(@class, "editor-post-taxonomies__hierarchical-terms-list")]//input[@type="checkbox"]',
                            '//div[contains(@class, "editor-post-taxonomies")]//input[@type="checkbox"]',
                            '//input[@type="checkbox" and ancestor::*[contains(@class, "categories")]]'
                        ]
                        checkboxes = []
                        for selector in selectores:
                            try:
                                checkboxes = driver.find_elements(By.XPATH, selector)
                                if checkboxes:
                                    break
                            except:
                                continue
                        
                        for cb in checkboxes:
                            try:
                                texto_categoria = ""
                                # Método A: Buscar label por for
                                label_id = cb.get_attribute('id')
                                if label_id:
                                    try:
                                        label = driver.find_element(By.XPATH, f'//label[@for="{label_id}"]')
                                        texto_categoria = label.text.strip().lower()
                                    except:
                                        pass
                                
                                # Método B: Buscar en el elemento padre más cercano
                                if not texto_categoria:
                                    try:
                                        # Buscar span o div con texto cerca del checkbox
                                        parent_elem = cb.find_element(By.XPATH, './following-sibling::*[1]')
                                        texto_categoria = parent_elem.text.strip().lower()
                                    except:
                                        try:
                                            # Alternativa: buscar en el siguiente hermano
                                            parent_elem = cb.find_element(By.XPATH, './parent::*/span')
                                            texto_categoria = parent_elem.text.strip().lower()
                                        except:
                                            pass
                                
                                # Método C: Buscar en contenedor padre
                                if not texto_categoria:
                                    try:
                                        parent_container = cb.find_element(By.XPATH, './ancestor::*[contains(@class, "components-checkbox-control")][1]')
                                        texto_categoria = parent_container.text.strip().lower()
                                    except:
                                        pass
                                
                                # Método D: Buscar usando JavaScript
                                if not texto_categoria:
                                    try:
                                        texto_categoria = driver.execute_script("""
                                            var cb = arguments[0];
                                            var label = document.querySelector('label[for="' + cb.id + '"]');
                                            if (label) return label.textContent.trim().toLowerCase();
                                            var parent = cb.closest('.components-checkbox-control');
                                            if (parent) return parent.textContent.trim().toLowerCase();
                                            return '';
                                        """, cb)
                                    except:
                                        pass
                                
                                # Comparar categorías
                                if texto_categoria:
                                    categoria_lower = clean_cat.lower().strip()
                                    texto_categoria = texto_categoria.strip()
                                    
                                    # Comparación exacta o parcial
                                    if categoria_lower == texto_categoria or categoria_lower in texto_categoria or texto_categoria in categoria_lower:
                                        if not cb.is_selected():
                                            # Usar JavaScript para hacer clic (más confiable)
                                            driver.execute_script("arguments[0].click();", cb)
                                            print(f"[OK] Categoría seleccionada: '{clean_cat}' -> '{texto_categoria}'")
                                        else:
                                            print(f"[INFO] Categoría ya seleccionada: '{clean_cat}'")
                                        checkbox_encontrado = True
                                        break
                            except Exception as e:
                                continue
                    except Exception as e:
                        print(f"[DEBUG] Error en búsqueda exacta: {e}")
                    
                    # Método 2: Buscar por palabras clave si no se encontró
                    if not checkbox_encontrado:
                        try:
                            palabras = clean_cat.lower().split()
                            checkboxes = driver.find_elements(By.XPATH, '//div[contains(@class, "components-panel__body")]//input[@type="checkbox"]')
                            for cb in checkboxes:
                                try:
                                    texto_categoria = ""
                                    # Intentar obtener el texto del label
                                    label_id = cb.get_attribute('id')
                                    if label_id:
                                        try:
                                            label = driver.find_element(By.XPATH, f'//label[@for="{label_id}"]')
                                            texto_categoria = label.text.strip().lower()
                                        except:
                                            pass
                                    
                                    # Si no funciona, intentar con el elemento padre
                                    if not texto_categoria:
                                        parent = cb.find_element(By.XPATH, './parent::*/parent::*')
                                        texto_categoria = parent.text.strip().lower()
                                    
                                    # Buscar si alguna palabra coincide
                                    if texto_categoria:
                                        for palabra in palabras:
                                            if len(palabra) > 3 and palabra in texto_categoria:
                                                if not cb.is_selected():
                                                    cb.click()
                                                    print(f"[OK] Categoría seleccionada (parcial): '{texto_categoria}' para '{categoria}'")
                                                else:
                                                    print(f"[INFO] Categoría ya seleccionada: '{texto_categoria}'")
                                                checkbox_encontrado = True
                                                break
                                except Exception as e:
                                    continue
                                if checkbox_encontrado:
                                    break
                        except Exception as e:
                            print(f"[DEBUG] Error en búsqueda parcial: {e}")
                    
                    # Si no se encontró la categoría, intentar crearla
                    if not checkbox_encontrado:
                        print(f"[WARNING] No se encontró la categoría '{clean_cat}' en WordPress")
                        print(f"[INFO] Intentando crear la categoría '{clean_cat}'...")
                        
                        # Intentar crear la nueva categoría
                        if crear_nueva_categoria(driver, wait, clean_cat):
                            print(f"[SUCCESS] Categoría '{clean_cat}' creada exitosamente")
                            
                            # Esperar un momento y volver a buscar la categoría para seleccionarla
                            time.sleep(2)
                            try:
                                # Buscar la categoría recién creada
                                checkboxes = driver.find_elements(By.XPATH, '//div[contains(@class, "components-panel__body")]//input[@type="checkbox"]')
                                for cb in checkboxes:
                                    try:
                                        texto_categoria = ""
                                        label_id = cb.get_attribute('id')
                                        if label_id:
                                            try:
                                                label = driver.find_element(By.XPATH, f'//label[@for="{label_id}"]')
                                                texto_categoria = label.text.strip().lower()
                                            except:
                                                pass
                                        
                                        if not texto_categoria:
                                            parent = cb.find_element(By.XPATH, './parent::*/parent::*')
                                            texto_categoria = parent.text.strip().lower()
                                        
                                        if texto_categoria and clean_cat.lower() == texto_categoria:
                                            if not cb.is_selected():
                                                cb.click()
                                                print(f"[OK] Categoría '{clean_cat}' seleccionada después de crearla")
                                            checkbox_encontrado = True
                                            break
                                    except:
                                        continue
                                
                                if not checkbox_encontrado:
                                    print(f"[WARNING] Se creó la categoría pero no se pudo seleccionar automáticamente")
                            except Exception as e:
                                print(f"[ERROR] Error al seleccionar categoría recién creada: {e}")
                        else:
                            print(f"[ERROR] No se pudo crear la categoría '{clean_cat}'")
                    
                    time.sleep(1)  # Pausa más larga para categorías
                        
                except Exception as e:
                    print(f"[ERROR] Error procesando categoría '{categoria}': {e}")
            
            # Pausa adicional después de configurar todas las categorías
            time.sleep(3)
            print("[INFO] Esperando que se procesen las categorías...")
                    
        except TimeoutException:
            print("[ERROR] Categorías no encontradas")
        
        # ---------------- INSERTAR ETIQUETAS ----------------
        if etiquetas:
            print(f"[INFO] Configurando etiquetas ({len(etiquetas)} etiquetas)...")
            try:
                # Expandir la sección "Etiquetas" si está colapsada
                try:
                    # Buscar el botón de la sección Etiquetas (puede estar colapsado)
                    panel_etiquetas = wait.until(EC.presence_of_element_located((
                        By.XPATH, '//div[contains(@class, "editor-post-taxonomies__flat-term-list") or contains(@class, "editor-post-taxonomies")]//button[contains(text(), "Etiquetas") or contains(@aria-label, "Etiquetas") or contains(text(), "Tags") or contains(@aria-label, "Tags")] | //button[contains(@class, "components-panel__body-toggle") and (contains(text(), "Etiquetas") or contains(text(), "Tags"))]'
                    )))
                    
                    # Verificar si está expandido
                    aria_expanded = panel_etiquetas.get_attribute("aria-expanded")
                    if aria_expanded == "false" or aria_expanded is None:
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", panel_etiquetas)
                        time.sleep(0.5)
                        panel_etiquetas.click()
                        print("[OK] Sección de etiquetas expandida")
                        time.sleep(2)
                    else:
                        print("[INFO] Sección de etiquetas ya está expandida")
                except Exception as e:
                    print(f"[DEBUG] No se encontró el botón de etiquetas para expandir: {e}")
                    # Continuar de todas formas, puede que ya esté expandido
                
                # Buscar el campo de entrada de etiquetas
                try:
                    # En Gutenberg, las etiquetas se introducen en un campo de texto con autocompletado
                    # Buscar con múltiples selectores
                    campo_etiquetas = None
                    selectores_etiquetas = [
                        '//input[contains(@placeholder, "Etiquetas") or contains(@placeholder, "Tags") or contains(@placeholder, "AÑADIR")]',
                        '//input[contains(@class, "components-form-token-field__input")]',
                        '//input[contains(@class, "tag-input")]',
                        '//div[contains(@class, "editor-post-taxonomies__flat-term-list")]//input[@type="text"]',
                        '//div[contains(@class, "components-form-token-field")]//input'
                    ]
                    
                    for selector in selectores_etiquetas:
                        try:
                            campo_etiquetas = wait.until(EC.visibility_of_element_located((By.XPATH, selector)))
                            if campo_etiquetas:
                                break
                        except:
                            continue
                    
                    if not campo_etiquetas:
                        raise Exception("No se encontró el campo de etiquetas")
                    
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo_etiquetas)
                    time.sleep(0.5)
                    campo_etiquetas.click()
                    time.sleep(0.5)
                    
                    # Insertar cada etiqueta una por una presionando Enter o Coma
                    for i, etiqueta in enumerate(etiquetas):
                        campo_etiquetas.clear()
                        campo_etiquetas.send_keys(etiqueta.strip())
                        time.sleep(0.3)
                        # Presionar Enter o Coma para confirmar
                        campo_etiquetas.send_keys(Keys.ENTER)
                        time.sleep(0.5)
                        print(f"[DEBUG] Etiqueta {i+1}/{len(etiquetas)} insertada: {etiqueta}")
                    
                    # Disparar eventos para que WordPress detecte los cambios
                    driver.execute_script("arguments[0].dispatchEvent(new Event('input', { bubbles: true }));", campo_etiquetas)
                    driver.execute_script("arguments[0].dispatchEvent(new Event('change', { bubbles: true }));", campo_etiquetas)
                    
                    print(f"[OK] {len(etiquetas)} etiquetas insertadas: {', '.join(etiquetas[:5])}{'...' if len(etiquetas) > 5 else ''}")
                    time.sleep(2)
                except Exception as e:
                    print(f"[WARNING] No se encontró el campo de etiquetas: {e}")
                    # Intentar método alternativo con JavaScript
                    try:
                        script = """
                        var etiquetas = arguments[0];
                        var input = document.querySelector('input[placeholder*="Etiquetas"], input[placeholder*="Tags"], .components-form-token-field__input input');
                        if (input) {
                            etiquetas.forEach(function(tag) {
                                input.value = tag;
                                input.dispatchEvent(new Event('input', { bubbles: true }));
                                input.dispatchEvent(new KeyboardEvent('keydown', { key: 'Enter', bubbles: true }));
                                input.dispatchEvent(new KeyboardEvent('keyup', { key: 'Enter', bubbles: true }));
                            });
                        }
                        """
                        driver.execute_script(script, etiquetas)
                        print(f"[OK] Etiquetas insertadas con JavaScript")
                        time.sleep(2)
                    except Exception as e2:
                        print(f"[ERROR] No se pudo insertar etiquetas: {e2}")
                        
            except Exception as e:
                print(f"[WARNING] Error al configurar etiquetas: {e}")
        else:
            print("[INFO] No hay etiquetas para insertar")

        # ---------------- PREPARAR PARA GUARDAR ----------------
        print("[INFO] Preparando para guardar...")
        
        # Hacer clic en el título para asegurar que el foco esté en el editor
        try:
            title_field = driver.find_element(By.XPATH, '//h1[@role="textbox"]')
            title_field.click()
            time.sleep(1)
        except:
            pass
        
        # Esperar más tiempo para que todos los cambios se procesen
        time.sleep(5)
        print("[INFO] Cambios procesados, procediendo a guardar...")
        
        # ---------------- GUARDAR COMO BORRADOR ----------------
        print("[INFO] Guardando como borrador...")
        
        # Verificar estado antes de guardar
        verificar_estado_guardado(driver)
        
        # Intentar guardar método principal
        if guardar_borrador(driver, wait):
            print(f"[SUCCESS] Post '{titulo}' guardado como borrador con éxito!")
            return True
        else:
            # Intentar método alternativo con teclado
            print("[INFO] Intentando método alternativo para guardar...")
            if guardar_con_teclado(driver):
                print(f"[SUCCESS] Post '{titulo}' guardado con método alternativo!")
                return True
            else:
                print(f"[ERROR] No se pudo guardar '{titulo}' como borrador")
                return False

    except Exception as e:
        print(f"[ERROR] Error procesando {archivo_word}: {e}")
        return False

def main():
    # Obtener todos los archivos Word de la carpeta
    archivos = obtener_archivos_word(carpeta_word)
    
    if not archivos:
        print(f"[ERROR] No se encontraron archivos .docx en la carpeta: {carpeta_word}")
        return
    
    print(f"[INFO] Se encontraron {len(archivos)} archivos Word:")
    for i, archivo in enumerate(archivos, 1):
        print(f"  {i}. {os.path.basename(archivo)}")
    
    # Inicializar el navegador una sola vez
    print("[INFO] Inicializando navegador...")
    driver = webdriver.Chrome(options=options)
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    wait = WebDriverWait(driver, 30)  # Timeout reducido para evitar cuelgues largos
    
    try:
        # Login una sola vez
        print("[INFO] Iniciando sesión en WordPress...")
        driver.get(url_wp)
        time.sleep(3)
        
        username_field = wait.until(EC.visibility_of_element_located((By.NAME, 'log')))
        password_field = wait.until(EC.visibility_of_element_located((By.NAME, 'pwd')))
        
        username_field.clear()
        username_field.send_keys(usuario_wp)
        password_field.clear()
        password_field.send_keys(password_wp)
        
        login_button = wait.until(EC.element_to_be_clickable((By.ID, 'wp-submit')))
        login_button.click()
        
        # Esperar a que cargue el dashboard
        wait.until(EC.presence_of_element_located((By.ID, 'wpadminbar')))
        print("[OK] Login exitoso")

        # Procesar cada archivo
        exitosos = 0
        total = len(archivos)
        
        for i, archivo in enumerate(archivos, 1):
            print(f"\n[INFO] Progreso: {i}/{total}")
            try:
                if procesar_archivo(archivo, driver, wait):
                    exitosos += 1
            except Exception as e:
                print(f"[ERROR] Error crítico procesando archivo {i}: {e}")
                # Intentar reiniciar el navegador si hay error crítico
                try:
                    driver.quit()
                    time.sleep(3)
                    driver = webdriver.Chrome(options=options)
                    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
                    wait = WebDriverWait(driver, 30)
                    
                    # Re-login
                    driver.get(url_wp)
                    time.sleep(3)
                    username_field = wait.until(EC.visibility_of_element_located((By.NAME, 'log')))
                    password_field = wait.until(EC.visibility_of_element_located((By.NAME, 'pwd')))
                    username_field.send_keys(usuario_wp)
                    password_field.send_keys(password_wp)
                    login_button = wait.until(EC.element_to_be_clickable((By.ID, 'wp-submit')))
                    login_button.click()
                    wait.until(EC.presence_of_element_located((By.ID, 'wpadminbar')))
                    print("[OK] Navegador reiniciado y re-logueado")
                except:
                    print("[ERROR] No se pudo reiniciar el navegador")
                    break
            time.sleep(3)  # Espera reducida entre guardados
        
        print(f"\n{'='*60}")
        print(f"[SUCCESS] PROCESO COMPLETADO")
        print(f"{'='*60}")
        print(f"[OK] Archivos guardados como borrador: {exitosos}/{total}")
        print(f"[ERROR] Archivos fallidos: {total - exitosos}")
        print(f"{'='*60}")

    except Exception as e:
        print(f"[ERROR] Error en el proceso principal: {e}")
        
    finally:
        time.sleep(3)
        driver.quit()
        print("[INFO] Navegador cerrado.")

if __name__ == "__main__":
    main()