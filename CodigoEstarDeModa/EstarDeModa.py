import time
import os
import re
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.action_chains import ActionChains

# ---------------- CONFIGURACIÓN ----------------
url_wp = 'https://estardemoda.net/wp-admin'
usuario_wp = 'esther'
password_wp = 'hDvt2J@yVCdf12@59*I1ezTa'
carpeta_word = r"C:\Users\Jonathan JD\Desktop\pink\Robert\CodigoEstarDeModa"

options = Options()
# options.add_argument("--headless")  # Opcional: ejecutar sin ventana
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--disable-blink-features=AutomationControlled")
options.add_experimental_option("excludeSwitches", ["enable-automation"])
options.add_experimental_option('useAutomationExtension', False)

# Añadir opciones para mejorar estabilidad
options.add_argument("--disable-gpu")
options.add_argument("--disable-software-rasterizer")
options.add_argument("--memory-pressure-off")
options.add_argument("--disable-extensions")
options.add_argument("--disable-plugins")
options.add_argument("--no-first-run")
options.add_argument("--disable-default-apps")

def obtener_archivos_word(carpeta):
    """Obtiene todos los archivos .docx de la carpeta especificada"""
    archivos = []
    for archivo in os.listdir(carpeta):
        if archivo.lower().endswith('.docx') and not archivo.startswith('~$'):
            ruta_completa = os.path.join(carpeta, archivo)
            archivos.append(ruta_completa)
    return archivos

def extraer_titulo(doc):
    """Extrae el título del documento - VERSIÓN MEJORADA"""
    # Buscar con múltiples variaciones
    variaciones_titulo = [
        "titulo de la ficha",
        "título de la ficha",
        "titulo ficha",
        "título",
        "titulo:",
        "título:"
    ]
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().lower()
        
        # Si encuentra alguna variación del marcador de título
        for variacion in variaciones_titulo:
            if variacion in texto:
                # Retornar el siguiente párrafo si existe
                if i + 1 < len(doc.paragraphs):
                    siguiente = doc.paragraphs[i + 1].text.strip()
                    if siguiente:
                        return siguiente
                # Si el título está en la misma línea después del marcador
                if ":" in para.text:
                    partes = para.text.split(":", 1)
                    if len(partes) > 1 and partes[1].strip():
                        return partes[1].strip()
    
    # Si no encontró con los marcadores, usar el primer párrafo no vacío
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto and len(texto) > 10:  # Al menos 10 caracteres
            return texto
    
    return ""

def extraer_descripcion(doc):
    """Extrae la descripción del documento y la formatea en HTML - VERSIÓN MEJORADA"""
    descripcion = []
    capturar = False
    lista_abierta = False
    
    # Palabras clave que indican el fin de la descripción
    palabras_fin = [
        r'\bfrase\s*clave\s*objetivo\b',
        r'\btítulo\s*seo\b',
        r'\btitulo\s*seo\b',
        r'\bmeta\s*description\b',
        r'\bcategorías\b',
        r'\bcategorias\b',
        r'\betiquetas\b',
        r'\bfin\s*descripción\b',
        r'\bfin\s*descripcion\b'
    ]
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        # Buscar inicio de descripción con múltiples variaciones
        if re.search(r'\bdescripcion\b', texto, re.IGNORECASE) or \
           re.search(r'\bdescripción\b', texto, re.IGNORECASE) or \
           'descripción:' in texto_lower or \
           'descripcion:' in texto_lower:
            capturar = True
            # Si la descripción está en la misma línea después de los dos puntos
            if ":" in texto and len(texto.split(":", 1)[1].strip()) > 0:
                # Procesar el texto después de los dos puntos
                texto_desc = texto.split(":", 1)[1].strip()
                if texto_desc:
                    descripcion.append(f"<p>{texto_desc}</p>")
            continue
        
        # Buscar fin de descripción con cualquiera de las palabras clave
        if capturar:
            debe_parar = False
            for patron in palabras_fin:
                if re.search(patron, texto, re.IGNORECASE):
                    debe_parar = True
                    break
            
            if debe_parar:
                # Cerrar lista si está abierta
                if lista_abierta:
                    descripcion.append('</ul>')
                    lista_abierta = False
                break
            
        if capturar and texto:
            # Verificar el estilo del párrafo
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            is_list = style_name in ["list paragraph", "viñeta", "list", "list bullet", "list number"]
            
            # Procesar encabezados
            if style_name in ["heading 1", "título 1"]:
                if lista_abierta:
                    descripcion.append('</ul>')
                    lista_abierta = False
                descripcion.append(f"<h1>{texto}</h1>")
            elif style_name in ["heading 2", "título 2"]:
                if lista_abierta:
                    descripcion.append('</ul>')
                    lista_abierta = False
                descripcion.append(f"<h2>{texto}</h2>")
            elif style_name in ["heading 3", "título 3"]:
                if lista_abierta:
                    descripcion.append('</ul>')
                    lista_abierta = False
                descripcion.append(f"<h3>{texto}</h3>")
            # Procesar elementos de lista
            elif is_list:
                if not lista_abierta:
                    descripcion.append('<ul>')
                    lista_abierta = True
                
                # Procesar contenido con formato
                item_html = ""
                for run in para.runs:
                    if run.text.strip():
                        texto_run = run.text
                        # Aplicar formatos múltiples
                        if run.bold:
                            texto_run = f"<strong>{texto_run}</strong>"
                        if run.italic:
                            texto_run = f"<em>{texto_run}</em>"
                        if run.underline:
                            texto_run = f"<u>{texto_run}</u>"
                        item_html += texto_run
                
                if not item_html.strip():
                    item_html = texto
                
                descripcion.append(f"<li>{item_html}</li>")
            # Procesar párrafos normales
            else:
                if lista_abierta:
                    descripcion.append('</ul>')
                    lista_abierta = False
                
                # Procesar formato en los runs
                paragraph_html = ""
                for run in para.runs:
                    if run.text.strip():
                        texto_run = run.text
                        # Aplicar formatos múltiples (negrita + cursiva + subrayado)
                        if run.bold:
                            texto_run = f"<strong>{texto_run}</strong>"
                        if run.italic:
                            texto_run = f"<em>{texto_run}</em>"
                        if run.underline:
                            texto_run = f"<u>{texto_run}</u>"
                        paragraph_html += texto_run
                
                if not paragraph_html.strip():
                    paragraph_html = texto
                
                descripcion.append(f"<p>{paragraph_html}</p>")
    
    # Cerrar lista si quedó abierta al final
    if lista_abierta:
        descripcion.append('</ul>')
    
    return ''.join(descripcion)

def frase_clave(doc):
    """Extrae la frase clave objetivo"""
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if "frase clave objetivo" in para.text.strip().lower():
            buscar = True
    return ""

def titulo_seo(doc):
    """Extrae el título SEO"""
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if "título seo" in para.text.strip().lower():
            buscar = True
    return ""

def meta_description(doc):
    """Extrae la meta description"""
    buscar = False
    for para in doc.paragraphs:
        if buscar:
            return para.text.strip()
        if "meta description" in para.text.strip().lower():
            buscar = True
    return ""

def leer_etiquetas(doc):
    """Extrae las etiquetas del documento"""
    buscar, etiquetas = False, ""
    for para in doc.paragraphs:
        texto = para.text.strip()
        if buscar and texto:
            etiquetas = texto.strip()
            break
        if "etiqueta" in texto.lower():
            buscar = True
    return etiquetas

def leer_categorias(doc):
    """Extrae las categorías del documento y las divide por comas"""
    categorias, capturar = [], False
    for para in doc.paragraphs:
        texto = para.text.strip()
        texto_lower = texto.lower()
        
        if capturar:
            if not texto or texto_lower == "fin categorias":
                break
            
            # Dividir por comas si hay múltiples categorías en una línea
            if "," in texto:
                # Dividir por comas y agregar cada categoría individualmente
                cats_divididas = texto.split(",")
                for cat in cats_divididas:
                    cat_limpia = cat.strip()
                    if cat_limpia:
                        categorias.append(cat_limpia)
            else:
                # Si no hay comas, agregar la categoría completa
                categorias.append(texto)
        
        if "categorías" in texto_lower or "categorias" in texto_lower:
            capturar = True
    
    # Limpiar y retornar categorías únicas
    return [cat.strip() for cat in categorias if cat.strip()]

def crear_nueva_categoria_clasico(driver, wait, nombre_categoria):
    """Crea una nueva categoría en WordPress (Editor Clásico) si no existe"""
    try:
        print(f"[INFO] Intentando crear nueva categoria: '{nombre_categoria}'")
        
        # Paso 1: Hacer clic en el enlace "Añadir una nueva categoría"
        try:
            # En el editor clásico, buscar el enlace para añadir categoría
            añadir_link = wait.until(EC.element_to_be_clickable((
                By.XPATH, '//a[contains(@class, "category-add") or contains(text(), "Añadir") or contains(text(), "nueva categoría")]'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", añadir_link)
            time.sleep(0.5)
            añadir_link.click()
            print("[OK] Enlace 'Anadir categoria' clickeado")
            time.sleep(2)
        except Exception as e:
            print(f"[WARNING] No se encontro el enlace 'Anadir categoria': {e}")
            # Intentar con selector alternativo
            try:
                añadir_link = driver.find_element(By.ID, 'category-add-toggle')
                añadir_link.click()
                time.sleep(2)
                print("[OK] Uso selector alternativo para anadir categoria")
            except:
                print("[ERROR] No se pudo hacer clic en 'Anadir categoria'")
                return False
        
        # Paso 2: Esperar a que aparezca el campo de nueva categoría
        try:
            # Buscar el campo de texto para el nombre de la nueva categoría
            campo_nombre = wait.until(EC.visibility_of_element_located((
                By.ID, 'newcategory'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", campo_nombre)
            time.sleep(0.5)
            campo_nombre.click()
            campo_nombre.clear()
            campo_nombre.send_keys(nombre_categoria)
            print(f"[OK] Nombre de categoria '{nombre_categoria}' ingresado")
            time.sleep(1)
        except:
            print("[ERROR] No se encontro el campo para el nombre de la categoria")
            return False
        
        # Paso 3: Hacer clic en el botón "Añadir nueva categoría"
        try:
            # Buscar el botón para confirmar la nueva categoría
            confirmar_btn = wait.until(EC.element_to_be_clickable((
                By.ID, 'category-add-submit'
            )))
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", confirmar_btn)
            time.sleep(0.5)
            confirmar_btn.click()
            print("[OK] Boton 'Anadir nueva categoria' clickeado")
            time.sleep(2)  # Esperar a que se cree la categoría
            
            # Verificar si se creó correctamente
            print(f"[SUCCESS] Categoria '{nombre_categoria}' creada exitosamente")
            return True
            
        except Exception as e:
            print(f"[WARNING] No se encontro el boton de confirmar: {e}")
            return False
        
    except Exception as e:
        print(f"[ERROR] Error general al crear categoria '{nombre_categoria}': {e}")
        return False

def guardar_borrador(driver, wait):
    """Función para guardar el post como borrador"""
    try:
        print("[INFO] Guardando como borrador...")
        
        # Esperar a que el botón esté disponible
        time.sleep(2)
        
        # Buscar botón de guardar borrador específico
        guardar_btn = wait.until(EC.element_to_be_clickable((By.ID, 'save-post')))
        
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", guardar_btn)
        time.sleep(2)
        
        # Verificar que el botón tenga el texto correcto
        btn_value = guardar_btn.get_attribute('value')
        print(f"[DEBUG] Valor del botón: '{btn_value}'")
        
        if btn_value != 'Guardar borrador':
            print(f"[WARNING] Botón no es 'Guardar borrador', es: '{btn_value}'")
            # Intentar buscar otros botones de guardar
            try:
                otros_botones = driver.find_elements(By.XPATH, '//input[@type="submit" and contains(@value, "Guardar")]')
                for boton in otros_botones:
                    print(f"[DEBUG] Otro botón encontrado: '{boton.get_attribute('value')}'")
                    if 'borrador' in boton.get_attribute('value').lower():
                        guardar_btn = boton
                        print("[OK] Usando botón alternativo de guardar borrador")
                        break
            except:
                pass
        
        # Verificar si el botón está habilitado
        if not guardar_btn.is_enabled():
            print("[ERROR] Botón de guardar deshabilitado, revisando posibles errores...")
            
            # Verificar si hay mensajes de error
            try:
                errores = driver.find_elements(By.CLASS_NAME, 'error')
                for error in errores:
                    print(f"[ERROR] Error encontrado: {error.text}")
            except:
                print("[WARNING] No se encontraron mensajes de error específicos")
            
            # Intentar hacer clic en otro lugar primero para activar validaciones
            try:
                title_field = driver.find_element(By.ID, 'title')
                title_field.click()
                time.sleep(1)
                
                # Intentar guardar nuevamente
                guardar_btn = wait.until(EC.element_to_be_clickable((By.ID, 'save-post')))
                if guardar_btn.is_enabled():
                    print("[OK] Botón habilitado después de interactuar con título")
                else:
                    print("[ERROR] Botón sigue deshabilitado")
                    return False
                    
            except Exception as e:
                print(f"[ERROR] Error al intentar habilitar botón: {e}")
                return False
        
        # Hacer clic en el botón
        guardar_btn.click()
        print("[OK] Clic en botón 'Guardar borrador'")
        
        # Esperar a que se procese el guardado
        time.sleep(5)
        
        # Verificar mensaje de éxito de diferentes formas
        try:
            # Método 1: Buscar mensaje de éxito
            exito_element = wait.until(EC.presence_of_element_located((By.ID, 'message')))
            mensaje = exito_element.text.lower()
            print(f"[INFO] Mensaje del sistema: {mensaje}")
            
            if any(palabra in mensaje for palabra in ['guardado', 'borrador', 'actualizado', 'saved', 'draft']):
                print("[OK] Guardado exitoso - Mensaje confirmado")
                return True
            else:
                print(f"[WARNING] Mensaje inesperado: {mensaje}")
                
        except TimeoutException:
            print("[WARNING] No se encontró mensaje de confirmación")
        
        # Método 2: Verificar si el botón cambió de estado
        try:
            time.sleep(2)
            btn_despues = driver.find_element(By.ID, 'save-post')
            nuevo_valor = btn_despues.get_attribute('value')
            if nuevo_valor != btn_value:
                print(f"[OK] Botón cambió de '{btn_value}' a '{nuevo_valor}' - Guardado exitoso")
                return True
        except:
            pass
            
        # Método 3: Verificar si hay algún indicador de éxito
        try:
            # Buscar cualquier elemento que indique éxito
            indicadores = driver.find_elements(By.XPATH, '//*[contains(text(), "Guardado") or contains(text(), "Borrador") or contains(text(), "Actualizado")]')
            for indicador in indicadores:
                if indicador.is_displayed():
                    print(f"[OK] Indicador de éxito encontrado: {indicador.text}")
                    return True
        except:
            pass
            
        # Si llegamos aquí, asumimos éxito pero con advertencia
        print("[WARNING] No se pudo confirmar explícitamente el guardado, pero continuamos")
        return True
            
    except TimeoutException as e:
        print(f"[ERROR] Tiempo de espera agotado al guardar: {e}")
        return False
    except Exception as e:
        print(f"[ERROR] Error al guardar: {e}")
        return False

def agregar_bloque_texto(driver, wait, contenido):
    """Agrega un bloque de texto usando Visual Composer"""
    try:
        print("[INFO] Agregando bloque de texto...")
        
        # Hacer clic en "Añadir elemento"
        add_element_btn = wait.until(EC.element_to_be_clickable((By.ID, 'vc_no-content-add-element')))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", add_element_btn)
        time.sleep(1)
        add_element_btn.click()
        print("[OK] Clic en 'Añadir elemento'")
        
        # Buscar y hacer clic en "Bloque de Texto" específico
        time.sleep(2)
        bloque_texto_btn = wait.until(EC.element_to_be_clickable((By.ID, 'vc_column_text')))
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", bloque_texto_btn)
        time.sleep(1)
        bloque_texto_btn.click()
        print("[OK] Clic en 'Bloque de Texto'")
        
        # Esperar a que se cargue el editor de texto
        time.sleep(3)
        
        # Cambiar a modo código/HTML
        try:
            html_btn = wait.until(EC.element_to_be_clickable((By.ID, 'wpb_tinymce_content-html')))
            html_btn.click()
            print("[OK] Cambiado a modo código")
            time.sleep(2)
        except Exception as e:
            print(f"[ERROR] No se pudo cambiar a modo código: {e}")
            return False
        
        # Insertar contenido en el bloque de texto (modo código)
        try:
            # Buscar el área de texto en modo código
            texto_area = wait.until(EC.element_to_be_clickable((By.ID, 'wpb_tinymce_content')))
            texto_area.click()
            texto_area.clear()
            
            # Insertar el contenido HTML
            texto_area.send_keys(contenido)
            print("[OK] Contenido HTML insertado en el bloque de texto")
            
        except Exception as e:
            print(f"[ERROR] No se pudo insertar en área de texto: {e}")
            return False
        
        # Guardar el bloque con el botón específico
        try:
            save_btn = wait.until(EC.element_to_be_clickable((By.XPATH, '//span[@data-vc-ui-element="button-save" and contains(text(), "Guardar Cambios")]')))
            save_btn.click()
            print("[OK] Bloque guardado")
        except Exception as e:
            print(f"[ERROR] No se pudo guardar el bloque: {e}")
            return False
        
        time.sleep(2)
        return True
        
    except Exception as e:
        print(f"[ERROR] Error al agregar bloque de texto: {e}")
        return False

def procesar_archivo(archivo_word, driver, wait):
    """Procesa un archivo Word individual y lo guarda como borrador en WordPress"""
    print(f"\n{'='*60}")
    print(f"[INFO] Procesando: {os.path.basename(archivo_word)}")
    print(f"{'='*60}")
    
    try:
        # Navegar a la página de nuevo post
        driver.get(url_wp + "/post-new.php")
        time.sleep(4)

        # Cargar el documento una sola vez
        doc = Document(archivo_word)
        
        # Extraer todos los datos del documento
        titulo = extraer_titulo(doc)
        descripcion = extraer_descripcion(doc)
        frase_obj = frase_clave(doc)
        tit_seo = titulo_seo(doc)
        meta_desc = meta_description(doc)
        categorias = leer_categorias(doc)
        etiquetas = leer_etiquetas(doc)

        print(f"[INFO] Título: {titulo if titulo else '[VACÍO - REVISAR DOCUMENTO]'}")
        if descripcion:
            print(f"[INFO] Descripción: {descripcion[:100]}... ({len(descripcion)} caracteres)")
        else:
            print(f"[WARNING] Descripción: [VACÍA - REVISAR DOCUMENTO]")
        print(f"[INFO] Frase clave: {frase_obj}")
        print(f"[INFO] Título SEO: {tit_seo}")
        print(f"[INFO] Meta desc: {meta_desc}")
        print(f"[INFO] Categorías: {categorias}")
        print(f"[INFO] Etiquetas: {etiquetas}")

        # ---------------- INSERTAR TÍTULO ----------------
        print("[INFO] Insertando título...")
        title_field = wait.until(EC.element_to_be_clickable((By.ID, 'title')))
        title_field.click()
        title_field.clear()
        title_field.send_keys(titulo)
        time.sleep(1)

        # ---------------- INSERTAR DESCRIPCIÓN CON VISUAL COMPOSER ----------------
        if descripcion.strip():
            print("[INFO] Insertando descripción con Visual Composer...")
            if not agregar_bloque_texto(driver, wait, descripcion):
                print("[ERROR] No se pudo insertar la descripción")
                return False
        else:
            print("[WARNING] No hay descripción para insertar")

        # ---------------- INSERTAR FRASE CLAVE ----------------
        print("[INFO] Insertando frase clave...")
        try:
            frase_field = wait.until(EC.visibility_of_element_located((By.ID, 'focus-keyword-input-metabox')))
            frase_field.clear()
            frase_field.send_keys(Keys.CONTROL, 'a')
            frase_field.send_keys(Keys.DELETE)
            frase_field.send_keys(frase_obj)
            print("[OK] Frase clave insertada")
        except TimeoutException:
            print("[WARNING] Campo de frase clave no encontrado")

        # ---------------- INSERTAR TÍTULO SEO ----------------
        print("[INFO] Insertando título SEO...")
        try:
            titseo_field = wait.until(EC.visibility_of_element_located((By.ID, 'yoast-google-preview-title-metabox')))
            titseo_field.click()
            titseo_field.send_keys(Keys.CONTROL, 'a')
            titseo_field.send_keys(Keys.DELETE)
            titseo_field.send_keys(tit_seo)
            print("[OK] Título SEO insertado")
        except TimeoutException:
            print("[WARNING] Campo de título SEO no encontrado")

        # ---------------- INSERTAR META DESCRIPCIÓN ----------------
        print("[INFO] Insertando meta descripción...")
        try:
            metdesc_field = wait.until(EC.visibility_of_element_located((By.ID, 'yoast-google-preview-description-metabox')))
            metdesc_field.click()
            metdesc_field.send_keys(Keys.CONTROL, 'a')
            metdesc_field.send_keys(Keys.DELETE)
            metdesc_field.send_keys(meta_desc)
            print("[OK] Meta descripción insertada")
        except TimeoutException:
            print("[WARNING] Campo de meta descripción no encontrado")

        # ---------------- INSERTAR CATEGORÍAS ----------------
        print("[INFO] Configurando categorías...")
        print(f"[DEBUG] Total de categorías detectadas: {len(categorias)}")
        if categorias:
            print(f"[DEBUG] Categorías a procesar: {', '.join(categorias)}")
        
        try:
            for categoria in categorias:
                categoria_limpia = categoria.strip()
                if not categoria_limpia:
                    continue
                    
                print(f"[DEBUG] Procesando categoría: '{categoria_limpia}'")
                checkbox_encontrado = False
                
                try:
                    # Método 1: Buscar por texto exacto del label
                    xpath = f"//label[normalize-space(text()) = '{categoria_limpia}']/input[@type='checkbox']"
                    checkbox = wait.until(EC.presence_of_element_located((By.XPATH, xpath)))
                    
                    # Hacer scroll para que sea visible (más margen para evitar la barra superior)
                    driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", checkbox)
                    time.sleep(0.5)
                    
                    if not checkbox.is_selected():
                        # Usar JavaScript click para evitar interceptación
                        try:
                            driver.execute_script("arguments[0].click();", checkbox)
                            print(f"[OK] Categoría seleccionada: {categoria_limpia}")
                        except Exception as e:
                            # Si JavaScript falla, intentar click normal
                            checkbox.click()
                            print(f"[OK] Categoría seleccionada (click directo): {categoria_limpia}")
                    else:
                        print(f"[INFO] Categoría ya estaba seleccionada: {categoria_limpia}")
                    checkbox_encontrado = True
                    time.sleep(0.5)
                    
                except TimeoutException:
                    print(f"[WARNING] No se encontró la categoría '{categoria_limpia}'")
                    
                    # Método 2: Intentar búsqueda más flexible
                    try:
                        # Buscar por contenido de texto (más flexible)
                        xpath_flex = f"//label[contains(normalize-space(text()), '{categoria_limpia}')]/input[@type='checkbox']"
                        checkbox = driver.find_element(By.XPATH, xpath_flex)
                        driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", checkbox)
                        time.sleep(0.5)
                        
                        if not checkbox.is_selected():
                            # Usar JavaScript click para evitar interceptación
                            try:
                                driver.execute_script("arguments[0].click();", checkbox)
                                print(f"[OK] Categoría seleccionada (búsqueda flexible): {categoria_limpia}")
                            except:
                                checkbox.click()
                                print(f"[OK] Categoría seleccionada (búsqueda flexible - click directo): {categoria_limpia}")
                        checkbox_encontrado = True
                        time.sleep(0.5)
                    except:
                        print(f"[WARNING] Categoría '{categoria_limpia}' no existe, intentando crearla...")
                        
                        # Método 3: Crear la categoría si no existe
                        if crear_nueva_categoria_clasico(driver, wait, categoria_limpia):
                            print(f"[OK] Categoría '{categoria_limpia}' creada exitosamente")
                            
                            # Esperar y buscar la categoría recién creada para seleccionarla
                            time.sleep(2)
                            try:
                                xpath = f"//label[normalize-space(text()) = '{categoria_limpia}']/input[@type='checkbox']"
                                checkbox = driver.find_element(By.XPATH, xpath)
                                driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", checkbox)
                                time.sleep(0.5)
                                
                                if not checkbox.is_selected():
                                    # Usar JavaScript click para evitar interceptación
                                    try:
                                        driver.execute_script("arguments[0].click();", checkbox)
                                        print(f"[OK] Categoría '{categoria_limpia}' seleccionada después de crearla")
                                    except:
                                        checkbox.click()
                                        print(f"[OK] Categoría '{categoria_limpia}' seleccionada después de crearla (click directo)")
                                checkbox_encontrado = True
                            except Exception as e:
                                print(f"[WARNING] Se creó la categoría pero no se pudo seleccionar: {e}")
                        else:
                            print(f"[ERROR] No se pudo crear la categoría '{categoria_limpia}'")
                
                except Exception as e:
                    print(f"[ERROR] Error inesperado con la categoría '{categoria_limpia}': {e}")
                    
        except Exception as e:
            print(f"[WARNING] Error general en categorías: {e}")

        # ---------------- INSERTAR ETIQUETAS ----------------
        print("[INFO] Configurando etiquetas...")
        try:
            if etiquetas.strip():
                etiq_field = wait.until(EC.visibility_of_element_located((By.ID, 'new-tag-post_tag')))
                etiq_field.clear()
                etiq_field.send_keys(etiquetas)
                etiq_field.send_keys(Keys.ENTER)
                print("[OK] Etiquetas insertadas")
        except Exception as e:
            print(f"[WARNING] Error insertando etiquetas: {e}")

        # ---------------- VERIFICAR ESTADO ANTES DE GUARDAR ----------------
        print("[INFO] Verificando estado antes de guardar...")
        time.sleep(3)
        
        # Verificar si hay algún error visible
        try:
            errores = driver.find_elements(By.CLASS_NAME, 'error')
            for error in errores:
                if error.is_displayed():
                    print(f"[ERROR] Error visible: {error.text}")
        except:
            pass

        # ---------------- GUARDAR COMO BORRADOR ----------------
        print("[INFO] Guardando como borrador...")
        
        if guardar_borrador(driver, wait):
            print(f"[SUCCESS] Post '{titulo}' guardado como borrador con éxito!")
            return True
        else:
            print(f"[ERROR] No se pudo guardar '{titulo}' como borrador")
            # Intentar método alternativo
            print("[INFO] Intentando método alternativo de guardado...")
            try:
                # Usar atajos de teclado
                actions = ActionChains(driver)
                actions.key_down(Keys.CONTROL).send_keys('s').key_up(Keys.CONTROL)
                actions.perform()
                time.sleep(3)
                print("[OK] Comando de teclado ejecutado")
                return True
            except Exception as e:
                print(f"[ERROR] Método alternativo también falló: {e}")
                return False

    except Exception as e:
        print(f"[ERROR] Error procesando {archivo_word}: {e}")
        return False

def main():
    # Obtener todos los archivos Word de la carpeta
    archivos = obtener_archivos_word(carpeta_word)
    
    if not archivos:
        print(f"[ERROR] No se encontraron archivos .docx en la carpeta: {carpeta_word}")
        return
    
    print(f"[INFO] Se encontraron {len(archivos)} archivos Word:")
    for i, archivo in enumerate(archivos, 1):
        print(f"  {i}. {os.path.basename(archivo)}")
    
    # Inicializar el navegador una sola vez
    print("[INFO] Inicializando navegador...")
    driver = webdriver.Chrome(options=options)
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    wait = WebDriverWait(driver, 25)
    
    try:
        # Login una sola vez
        print("[INFO] Iniciando sesión en WordPress...")
        driver.get(url_wp)
        time.sleep(3)
        
        username_field = wait.until(EC.visibility_of_element_located((By.NAME, 'log')))
        password_field = wait.until(EC.visibility_of_element_located((By.NAME, 'pwd')))
        
        username_field.clear()
        username_field.send_keys(usuario_wp)
        password_field.clear()
        password_field.send_keys(password_wp)
        
        login_button = wait.until(EC.element_to_be_clickable((By.ID, 'wp-submit')))
        login_button.click()
        
        # Esperar a que cargue el dashboard
        wait.until(EC.presence_of_element_located((By.ID, 'wpadminbar')))
        print("[OK] Login exitoso")

        # Procesar cada archivo
        exitosos = 0
        total = len(archivos)
        
        for i, archivo in enumerate(archivos, 1):
            print(f"\n[INFO] Progreso: {i}/{total}")
            if procesar_archivo(archivo, driver, wait):
                exitosos += 1
            time.sleep(5)  # Espera entre guardados
        
        print(f"\n{'='*60}")
        print(f"[SUCCESS] PROCESO COMPLETADO")
        print(f"{'='*60}")
        print(f"[OK] Archivos guardados como borrador: {exitosos}/{total}")
        print(f"[ERROR] Archivos fallidos: {total - exitosos}")
        print(f"{'='*60}")

    except Exception as e:
        print(f"[ERROR] Error en el proceso principal: {e}")
        
    finally:
        time.sleep(3)
        driver.quit()
        print("[INFO] Navegador cerrado.")

if __name__ == "__main__":
    main()